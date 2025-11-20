from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docxtpl import DocxTemplate
from datetime import datetime
import os
from pathlib import Path

context = {}


def arrest_aftoforo():
    global context
    doc = Document()
    title = doc.add_heading('ΕΚΘΕΣΗ ΣΥΛΛΗΨΗΣ ΚΑΙ ΒΕΒΑΙΩΣΗΣ ΠΛΗΜΜΕΛΗΜΑΤΟΣ ΕΠ\' ΑΥΤΟΦΩΡΟ(Κ.Π.Δ)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph = doc.add_paragraph()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run(
        '    Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}} ημέρα εβδομάδας '
        '{{day}} και ώρα {{hour}} ο κάτωθι υπογεγραμμένος  {{officer}} του {{policeStation}} '
        'Θεσσαλονίκης, εκθέτω τα παρακάτω: ' )

    new_para1 = doc.add_paragraph(' Την {{hourOfCrime}} ώρα στις {{dateOfCrime}} κατέλαβα επ’ αυτοφώρω στη '
                                  '{{placeOfCrime}} τον {{surnamePerperator}} {{namePerperator}} του '
                                  "{{fathernamePerperator}} και της {{mothernamePerperator}}"
                                  "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ."
                                  " {{addressPerperator}},αριθμός τηλεφώνου {{telPreperator}}, ηλεκτρονικό "
                                  " ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}} που "
                                  "εκδόθηκε την {{issuedPerperator}} από {{place_issuedPerperator}} "
                                  " Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}}, να λειτουργεί ως προσωρινά "
                                  "υπεύθυνος κατάστημα υγειονομικού ενδιαφέροντος, ήτοι {{typeOfStore}}  με διακριτικό"
                                  "τίτλο {{nameOfStore}} ιδιοκτησίας του {{surname}} {{name}} του {{fathername}} "
                                  "και της"
                                  ' {{mothername}} γεν. {{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός '
                                  'τηλεφώνου {{tel}}, ηλεκτρονικό ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DAT}}'
                                  ' που εκδόθηκε την {{issued}} από {{place_issued}} Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}} '
                                  'ο οποίος απουσίαζε κατά την ώρα του ελέγχου, καταμηνύεται με την παρούσα και '
                                  'αναζητείται στα όρια του αυτοφώρου και να διαταράσσει την κοινή ησυχία των περιοίκων'
                                  ' με μουσική προερχόμενη από στερεοφωνικό συγκρότημα και ενισχυτή. Διευκρινίζεται ότι'
                                  'η μέτρηση έγινε με ειδικό ηχόμετρο ρυθμισμένο να μετρά την Α’ ηχοστάθμη και η '
                                  'μουσική βρέθηκε να έχει μέγιστη επαναλαμβανόμενη ηχοστάθμη(ΜΕΗ) {{measurement}} '
                                  '{{measurNum}} dB αντί του ανώτατου επιτρεπόμενου ορίου ογδόντα (80) dB όπως '
                                  'φαίνεται και στο συνημμένο φύλλο ελέγχου θορύβου κατά παράβαση του άρθρου 14 '
                                  'παράγραφος 4 της υπ’ αριθμόν Υ1γ/ΓΠ/47829/17ΥΔ και Α5/3010/85 Υ.Δ όπως τροποποιήθηκε'
                                  ' με Υ2/ΟΙΚ 15438/2001 Κ.Υ.Α σε συνδυασμό με άρθρο 37 παράγραφος 3 του Ν.4055/2012.')

    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph(
                                'Κατόπιν των ανωτέρω ενημέρωσα αυτόν για τα προβλεπόμενα στο άρθρο 96 Κ.Π.Δ δικαιώματα '
                                'του, του παρείχα το προβλεπόμενο στην παράγραφο 1 του άρθρου αυτού έγγραφο στην '
                                'κατάλληλη γλώσσα και πείσθηκα ότι ουδεμία αμφιβολία υπάρχει ως προς την ταυτότητα του,'
                                'συνέλαβα αυτόν για να τον προσάγω στον Εισαγγελέα Πλημμελειοδικών Θεσσαλονίκης'
                                ' για άμεση παραπομπή του στο ακροατήριο του Πλημμελειοδικείου, σύμφωνα με τα άρθρα 417'
                                ' – 426 του Κώδικα Ποινικής Δικονομίας, επειδή το αδίκημα που διαπράχθηκε υπάγεται στη '
                                'συνοπτική διαδικασία των άρθρων αυτών.')
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    new_para2 = doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hour}} ώρα και περαιώθηκε την "
                                  "{{hourOfReportFinished}} ώρα. Για πιστοποίηση συντάχθηκε η παρούσα έκθεση η οποία "
                                  "αφού"
                                  "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para3 = doc.add_paragraph("                                                                                   "
                                  "                          Ο   Συλλαβών και Βεβαιών την πράξη ")
    new_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def create_desktop_folder(folder_name):
        """Δημιουργία φακέλου στην επιφάνεια εργασίας"""
        # Βρες το Desktop path
        desktop = Path.home() / "Desktop"
        # Δημιούργησε τον φάκελο
        folder_path = desktop / folder_name
        folder_path.mkdir(exist_ok = True)  # exist_ok=True δεν δίνει error αν υπάρχει ήδη
        return folder_path

    folder_name = "Police_Reports"
    folder_path = create_desktop_folder(folder_name)
    file_path = folder_path / "Εκθεση_Πλημ_Επ_Αυτοφ.docx"

    template = DocxTemplate("Εκθεση_Πλημ_Επ_Αυτοφ.docx")
    template.render(context)
    template.save(str(file_path))



def arrest_aftoforo2():
    global context
    doc = Document()
    title = doc.add_heading('ΕΚΘΕΣΗ ΣΥΛΛΗΨΗΣ ΚΑΙ ΒΕΒΑΙΩΣΗΣ ΠΛΗΜΜΕΛΗΜΑΤΟΣ ΕΠ\' ΑΥΤΟΦΩΡΟ(Κ.Π.Δ)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph = doc.add_paragraph()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run(
        '    Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}} ημέρα εβδομάδας '
        '{{day}} και ώρα {{hour}} ο κάτωθι υπογεγραμμένος  {{officer}} του {{policeStation}} '
        'Θεσσαλονίκης, εκθέτω τα παρακάτω: ' )

    new_para1 = doc.add_paragraph(' Την {{hourOfCrime}} ώρα στις {{dateOfCrime}} κατέλαβα επ’ αυτοφώρω στη '
                                  '{{placeOfCrime}} τον {{surnamePerperator}} {{namePerperator}} του '
                                  "{{fathernamePerperator}} και της {{mothernamePerperator}}"
                                  "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ."
                                  " {{addressPerperator}},αριθμός τηλεφώνου {{telPreperator}}, ηλεκτρονικό "
                                  " ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}} που "
                                  "εκδόθηκε την {{issuedPerperator}} από {{place_issuedPerperator}} "
                                  " Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}}, να λειτουργεί ως ιδιοκτήτης"
                                  " κατάστημα υγειονομικού ενδιαφέροντος, ήτοι {{typeOfStore}}  με διακριτικό"
                                  'τίτλο {{nameOfStore}}  να διαταράσσει την κοινή ησυχία των περιοίκων'
                                  ' με μουσική προερχόμενη από στερεοφωνικό συγκρότημα και ενισχυτή. Διευκρινίζεται ότι'
                                  'η μέτρηση έγινε με ειδικό ηχόμετρο ρυθμισμένο να μετρά την Α’ ηχοστάθμη και η '
                                  'μουσική βρέθηκε να έχει μέγιστη επαναλαμβανόμενη ηχοστάθμη(ΜΕΗ) {{measurement}} '
                                  '{{measurNum}} dB αντί του ανώτατου επιτρεπόμενου ορίου ογδόντα (80) dB όπως '
                                  'φαίνεται και στο συνημμένο φύλλο ελέγχου θορύβου κατά παράβαση του άρθρου 14 '
                                  'παράγραφος 4 της υπ’ αριθμόν Υ1γ/ΓΠ/47829/17ΥΔ και Α5/3010/85 Υ.Δ όπως τροποποιήθηκε'
                                  ' με Υ2/ΟΙΚ 15438/2001 Κ.Υ.Α σε συνδυασμό με άρθρο 37 παράγραφος 3 του Ν.4055/2012.')

    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph(
                                'Κατόπιν των ανωτέρω ενημέρωσα αυτόν για τα προβλεπόμενα στο άρθρο 96 Κ.Π.Δ δικαιώματα '
                                'του, του παρείχα το προβλεπόμενο στην παράγραφο 1 του άρθρου αυτού έγγραφο στην '
                                'κατάλληλη γλώσσα και πείσθηκα ότι ουδεμία αμφιβολία υπάρχει ως προς την ταυτότητα του,'
                                'συνέλαβα αυτόν για να τον προσάγω στον Εισαγγελέα Πλημμελειοδικών Θεσσαλονίκης'
                                ' για άμεση παραπομπή του στο ακροατήριο του Πλημμελειοδικείου, σύμφωνα με τα άρθρα 417'
                                ' – 426 του Κώδικα Ποινικής Δικονομίας, επειδή το αδίκημα που διαπράχθηκε υπάγεται στη '
                                'συνοπτική διαδικασία των άρθρων αυτών.')
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    new_para2 = doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hour}} ώρα και περαιώθηκε την "
                                  "{{hourOfReportFinished}} ώρα. Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού"
                                  "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para3 = doc.add_paragraph("                                                                                   "
                                  "         Ο   Συλλαβών και Βεβαιών την πράξη ")
    new_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('Εκθεση_Πλημ_Επ_Αυτοφ2.docx')
    template = DocxTemplate("Εκθεση_Πλημ_Επ_Αυτοφ2.docx")

    template.render(context)
    template.save("Εκθεση_Πλημ_Επ_Αυτοφ2.docx")


def seizure():
    global context
    doc=Document()
    title=doc.add_heading('ΕΚΘΕΣΗ ΠΑΡΑΔΟΣΕΩΣ ΚΑΙ ΚΑΤΑΣΧΕΣΕΩΣ (Κ.Π.Δ)', level = 0)
    title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph=doc.add_paragraph()
    main_paragraph.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run( 'Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}}'
                            ' ημέρα εβδομάδας {{day}} και ώρα {{hourSeizure}} ενώπιον εμού του {{first_officer}} '
                            'του {{policeStation}} Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της '
                            'ιδίας υπηρεσίας, που προσλήφθηκε ως Β\' Ανακριτικός Υπάλληλος, εκθέτουμε τα ακόλουθα:'
                            'επειδή ενεργούμε προανάκριση για παράβαση του άρθρου 14 παράγραφος 4 της υπ’ αριθμόν'
                            'Υ1γ/ΓΠ/47829/17ΥΔ και Α5/3010/85 Υ.Δ όπως τροποποιήθηκε με Υ2/ΟΙΚ 15438/2001 Κ.Υ.Α σε'
                            'συνδυασμό με άρθρο 37 παράγραφος 3 του Ν.4055/2012, πράξη που έλαβε χώρα την '
                            '{{hourOfCrime}} ώρα της {{dateOfCrime}} στο κατάστημα υγειονομικού ενδιαφέροντος, ήτοι'
                            '{{typeOfStore}} {{nameOfStore}} με διεύθυνση {{placeOfCrime}}, ιδιοκτησίας {{surname}}'
                            ' {{name}} του {{fathername}} και της {{mothername}} γεν. {{dateOfBirth}} στη '
                            '{{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου {{tel}}, ηλεκτρονικό ταχυδρομείο '
                            '{{email}}, κάτοχος του υπ αριθμόν {{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}}'
                            ' Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}}, προβήκαμε στην κατάσχεση ενός {{seizure}} ο οποίος μας '
                            'παραδόθηκε από τον προσωρινά υπεύθυνο του εν λόγω καταστήματος {{surnamePerperator}} '
                            '{{namePerperator}} του {{fathernamePerperator}} και της {{mothernamePerperator}}'
                            "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ."
                            " {{addressPerperator}},αριθμός τηλεφώνου {{telPreperator}}, ηλεκτρονικό "
                            " ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}} που "
                            "εκδόθηκε την {{issuedPerperator}} από {{place_issuedPerperator}} "
                            ' Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}}.'

                            )
    new_para1=doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hourSeizure}} ώρα και περαιώθηκε την "
                                "{{hourOfSeizureFinished}} ώρα. Για πίστοποίηση συντάχθηκε η παρούσα έκθεση η οποία "
                                "αφού αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5=doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sample_document.docx')
    template=DocxTemplate("sample_document.docx")
    template.render(context)
    template.save("Εκθεση_Παράδοσης_Κατασχεσης.docx")


def seizure2():
    global context
    doc=Document()
    title=doc.add_heading('ΕΚΘΕΣΗ ΠΑΡΑΔΟΣΕΩΣ ΚΑΙ ΚΑΤΑΣΧΕΣΕΩΣ (Κ.Π.Δ)', level = 0)
    title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph=doc.add_paragraph()
    main_paragraph.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run( 'Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}}'
                            ' ημέρα εβδομάδας {{day}} και ώρα {{hourSeizure}} ενώπιον εμού του {{first_officer}} '
                            'του {{policeStation}} Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της '
                            'ιδίας υπηρεσίας, που προσλήφθηκε ως Β\' Ανακριτικός Υπάλληλος, εκθέτουμε τα ακόλουθα:'
                            'επειδή ενεργούμε προανάκριση για παράβαση του άρθρου 14 παράγραφος 4 της υπ’ αριθμόν'
                            'Υ1γ/ΓΠ/47829/17ΥΔ και Α5/3010/85 Υ.Δ όπως τροποποιήθηκε με Υ2/ΟΙΚ 15438/2001 Κ.Υ.Α σε'
                            'συνδυασμό με άρθρο 37 παράγραφος 3 του Ν.4055/2012, πράξη που έλαβε χώρα την '
                            '{{hourOfCrime}} ώρα της {{dateOfCrime}} στο κατάστημα υγειονομικού ενδιαφέροντος, ήτοι'
                            '{{typeOfStore}} {{nameOfStore}} με διεύθυνση {{placeOfCrime}}, ιδιοκτησίας {{surname}}'
                            ' {{name}} του {{fathername}} και της {{mothername}} γεν. {{dateOfBirth}} στη '
                            '{{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου {{tel}}, ηλεκτρονικό ταχυδρομείο '
                            '{{email}}, κάτοχος του υπ αριθμόν {{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}}'
                            ' Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}}, προβήκαμε στην κατάσχεση ενός {{seizure}} ο οποίος μας '
                            'παραδόθηκε από τον ανωτέρω ιδιοκτήτη του εν λόγω καταστήματος.'

                            )
    new_para1=doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hourSeizure}} ώρα και περαιώθηκε την "
                                "{{hourOfSeizureFinished}} ώρα. Για πίστοποίηση συντάχθηκε η παρούσα έκθεση η οποία "
                                "αφού αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5=doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sample_document.docx')
    template=DocxTemplate("sample_document.docx")
    template.render(context)
    template.save("Εκθεση_Παράδοσης_Κατασχεσης.docx")

def deltio_kat2():
    global context

    # Δημιουργία εγγράφου
    doc = Document ( )
    title = doc.add_heading ( 'ΔΕΛΤΙΟ ΣΤΟΙΧΕΙΩΝ ΤΑΥΤΟΤΗΤΑΣ ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ' , level = 0 )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ορισμός γραμματοσειράς (bold + υπογραμμισμένο) για τον τίτλο
    for run in title.runs:
        run.font.size = Pt ( 14 )
        run.bold = True  # Bold
        run.underline = WD_UNDERLINE.SINGLE  # Υπογράμμιση

        # ΕΠΩΝΥΜΟ
        p_eponymo = doc.add_paragraph ( )
        run_label = p_eponymo.add_run ( "ΕΠΩΝΥΜΟ: " )
        run_label.bold = True
        run_value = p_eponymo.add_run ( " {{ surname }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝΟΜΑ
        p_onoma = doc.add_paragraph ( )
        run_label = p_onoma.add_run ( "ΟΝΟΜΑ:    " )
        run_label.bold = True
        run_value = p_onoma.add_run ( "{{ name }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝΟΜΑ ΠΑΤΕΡΑ
        p_onoma_patera = doc.add_paragraph ( )
        run_label = p_onoma_patera.add_run ( "ΟΝΟΜΑ ΠΑΤΕΡΑ:    " )
        run_label.bold = True
        run_value = p_onoma_patera.add_run ( "{{ fathername }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝΟΜΑ ΜΗΤΕΡΑΣ
        p_onoma_mhteras = doc.add_paragraph ( )
        run_label = p_onoma_mhteras.add_run ( "ΟΝΟΜΑ ΜΗΤΕΡΑΣ:    " )
        run_label.bold = True
        run_value = p_onoma_mhteras.add_run ( " {{ mothername }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΠΑΤΡΙΚΟ ΕΠΩΝΥΜΟ (Για έγγαμες)
        p_onoma_patriko = doc.add_paragraph ( )
        run_label = p_onoma_patriko.add_run ( "ΠΑΤΡΙΚΟ ΕΠΩΝΥΜΟ (Για έγγαμεσ):    " )
        run_label.bold = True
        run_value = p_onoma_patriko.add_run ( "_______________________" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝ/ΜΟ ΣΥΖΥΓΟΥ
        p_onoma_syzygou = doc.add_paragraph ( )
        run_label = p_onoma_syzygou.add_run ( "ΟΝ/ΜΟ ΣΥΖΥΓΟΥ:    " )
        run_label.bold = True
        run_value = p_onoma_syzygou.add_run ( "_______________________" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΗΜΕΡΟΜΗΝΙΑ ΓΕΝΝΗΣΗΣ
        p_dateofbirth = doc.add_paragraph ( )
        run_label = p_dateofbirth.add_run ( "ΗΜΕΡΟΜΗΝΙΑ ΓΕΝΝΗΣΗΣ:    " )
        run_label.bold = True
        run_value = p_dateofbirth.add_run ( "  {{ dateOfBirth  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΤΟΠΟΘΕΣΙΑ ΓΕΝΝΗΣΗΣ
        p_placeofbirth = doc.add_paragraph ( )
        run_label = p_placeofbirth.add_run ( " ΤΟΠΟΘΕΣΙΑ ΓΕΝΝΗΣΗΣ:    " )
        run_label.bold = True
        run_value = p_placeofbirth.add_run ( "  {{ placeOfBirth  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΤΟΠΟΘΕΣΙΑ κατοικιασ
        p_address = doc.add_paragraph ( )
        run_label = p_address.add_run ( " ΤΟΠΟΣ ΚΑΤΟΙΚΙΑΣ:    " )
        run_label.bold = True
        run_value = p_address.add_run ( "  {{ address  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # Δ.Α.Τ
        p_dat = doc.add_paragraph ( )
        run_label = p_dat.add_run ( " ΑΡΙΘΜΟΣ Δ.Α.Τ:    " )
        run_label.bold = True
        run_value = p_dat.add_run ( "  {{ DAT  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # Δ.Α.Τ
        p_issue = doc.add_paragraph ( )
        run_label = p_issue.add_run ( " ΕΚΔΟΘΕΝ:    " )
        run_label.bold = True
        run_value = p_issue.add_run ( "  {{ issuedPerperator  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # Δ.Α.Τ
        p_issueplace = doc.add_paragraph ( )
        run_label = p_issueplace.add_run ( " ΑΠΟ:    " )
        run_label.bold = True
        run_value = p_issueplace.add_run ( " {{ place_issued  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΑΦΜ
        # Δ.Α.Τ
        p_afm = doc.add_paragraph ( )
        run_label = p_afm.add_run ( " Α.Φ.Μ:    " )
        run_label.bold = True
        run_value = p_afm.add_run ( "  {{ afm }} από Δ.Ο.Υ {{  doy }}  " )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        footer_para = doc.add_paragraph ( )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Προσθήκη κειμένου με ακριβή θέση (χρήση tabs)
        footer_run = footer_para.add_run ( )
        footer_run.add_text ( f"\t\tΘέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
        footer_run.add_text ( "\t\t\tΟ\t\t\n" )  # 5 tabs για στοίχιση
        footer_run.add_text ( "\t\t\t\t\tΑνακριτικός Υπάλληλος\n\n" )
        footer_run.add_text ( "\t\t\t\t\t {{ first_officer }}\n" )

        # Αποθήκευση
        doc.save ( 'deltio_katigoroumenoutemplate2.docx' )
        # Επεξεργασία και τελική αποθήκευση
        template = DocxTemplate ( 'deltio_katigoroumenoutemplate2.docx' )
        template.render ( context )
        template.save ( "ΔΕΛΤΙΟ_ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ2.docx" )
        # Αποθήκευση του εγγράφου




def transmission_sound():
    global context  # Declare we're using the global context
    # Δημιουργία εγγράφου
    doc = Document ( )

    # Δημιουργία πίνακα με 1 γραμμή και 2 κελιά (για αριστερά & δεξιά κείμενο)
    table = doc.add_table ( rows = 1 ,cols = 2 )
    table.autofit = False  # Απενεργοποίηση αυτόματου προσαρμογής πλάτους

    # Ορισμός πλάτους στηλών (50% - 50%)
    col_widths = [Inches ( 3 ) ,Inches ( 3 )]
    for i ,width in enumerate ( col_widths ):
        table.columns[i].width = width

    # Κελί 1 (Αριστερά) - "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ"
    left_cell = table.cell ( 0 ,0 )
    left_texts = [
        "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ" ,
        "ΕΛΛΗΝΙΚΗ ΑΣΤΥΝΟΜΙΑ" ,
        "ΔΙΕΥΘ. ΑΣΤ. ΘΕΣΣΑΛΟΝΙΚΗΣ" ,
        "ΑΣΤΥΝΟΜΙΚΟ ΤΜΗΜΑ ΘΕΡΜΗΣ" ,
        "ΤΗΛ:2310461203" ,
        "email: atthermis@astynomia.gr" ,
        "Αρμόδιος:{{first_officer }}" ,
        "Αρ.πρωτ: {{ protocolnumber }}"
    ]

    for text in left_texts:
        p = left_cell.add_paragraph ( )
        run = p.add_run ( text )
        run.bold = True
        run.font.size = Pt ( 10 )  # γραμματοσειρά (12pt)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt ( 4 )

    # Κελί 2 (Δεξιά) - Τοποθεσία, ημερομηνία + επιπλέον κείμενο
    right_cell = table.cell ( 0 ,1 )

    # Γραμμή 1: Τοποθεσία & Ημερομηνία
    right_para1 = right_cell.add_paragraph ( )
    right_run1 = right_para1.add_run ( f"Θέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
    right_run1.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Γραμμή 2: "Π Ρ Ο Σ" (με κενά)
    right_para2 = right_cell.add_paragraph ( )
    right_run2 = right_para2.add_run ( " Π Ρ Ο Σ" )
    right_run2.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Γραμμή 3: Εργαστήριο & Πανεπιστήμιο
    right_para3 = right_cell.add_paragraph ( )
    right_run3 = right_para3.add_run (
        "ΕΙΣΑΓΓΕΛΙΑ ΠΛΗΜ/ΚΩΝ \nΘΕΣΣΑΛΟΝΙΚΗΣ\n" )
    right_run3.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph ( ).add_run ( ).add_break ( )  # Κενή γραμμή πριν το ΘΕΜΑ

    p_theme = doc.add_paragraph ( )
    run_theme = p_theme.add_run ( "ΘΕΜΑ: «Αποστολή Συνοδείας»" )
    run_theme.bold = True
    run_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    p_theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν

    main_theme = doc.add_paragraph()
    run1_theme = main_theme.add_run(  '   Αποστέλλεται συνοδεία Αστυνομικών Υπηρεσίας μας  και και με την σε βάρος του'
                                      ' σχηματισθείσα Δικογραφία ο {{surnamePerperator}} '
                                      '{{namePerperator}} του {{fathernamePerperator}} και της {{mothernamePerperator}}'
                                      'γεν. {{dateOfBirthPerperator }} στη {{placeOfBirthPerperator }} κατ.{'
                                      '{ addressPerperator }},αριθμός τηλεφώνου {{ telPreperator }}, ηλεκτρονικό '
                                      'ταχυδρομείου{{ emailPreperator }}, κάτοχος του υπ αριθμόν {{ DATperperator}} που'
                                      ' εκδόθηκε την {{ issuedPerperator }} από {{place_issuedPerperator}} Α.Φ.Μ : '
                                      '{{afmPreperator}}, Δ.Ο.Υ : {{ doyPrep }}, κατηγορούμενος παράβαση του άρθρου 14'
                                      'παρ.4 της υπ’ αριθμόν Υ1γ/ΓΠ/47829/17ΥΔ και Α5/3010/85 Υ.Δ όπως τροποποιήθηκε με'
                                      'Υ2/ΟΙΚ 15438/2001 Κ.Υ.Α σε συνδυασμό με άρθρο 37 παράγραφος 3 του Ν.4055/2012.'
                                      )
    run1_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    new_para1 = doc.add_paragraph('- Ειδικότερα, την {{hourOfCrime}} ώρα της {{dateOfCrime}} συνελήφθη από Αστυνομικό'
                                  'του {{policeStation2}} στην οδό {{placeOfCrime}}, όπου εδρεύει το κατάστημα '
                                  'υγειονομικού ενδιαφέροντος {{typeOfStore}} με την επωνυμία {{nameOfStore}} '
                                  'ιδιοκτησίας {{surname}} {{name}} του {{fathername}} και της {{mothername}} γεν.'
                                  ' {{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου {{tel}}, '
                                  'ηλεκτρονικό ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DAT}} που εκδόθηκε την '
                                  '{{issued}} από {{place_issued}} να διαταράσσει την κοινή ησυχία των περιοίκων με '
                                  'μουσική προερχόμενη από στερεοφωνικό συγκρότημα και ενισχυτή. Διευκρινίζεται ότι η'
                                  ' μέτρηση έγινε με ειδικό ηχόμετρο ρυθμισμένο να μετρά την Α’ ηχοστάθμη και η μουσική'
                                  'βρέθηκε να έχει μέγιστη επαναλαμβανόμενη ηχοστάθμη(ΜΕΗ) {{measurement}}{{measurNum}}'
                                  'dB αντί του ανώτατου επιτρεπόμενου ορίου ογδόντα (80) dB όπως φαίνεται και στο'
                                  ' συνημμένο φύλλο ελέγχου θορύβου.')
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph('Ανωτέρω ιδιοκτήτης κατά τον έλεγχο δεν βρισκόταν στο σημείο και αναζητείται εντός '
                                  'ορίων αυτοφώρου.')

    footer_para = doc.add_paragraph ( )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Προσθήκη κειμένου με ακριβή θέση (χρήση tabs)
    footer_run = footer_para.add_run ( )
    footer_run.add_text(f"\t\tΘέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
    footer_run.add_text("\t\t\tΟ\t\t\n")  # 5 tabs για στοίχιση
    footer_run.add_text("\t\t\t\t\tΔΙΟΙΚΗΤΗΣ\n\n")


    # Αποθήκευση
    doc.save('sample_transmission1.docx')
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ('sample_transmission1.docx')
    template.render(context)
    template.save("ΔΙΑΒΙΒΑΣΤΙΚΟ_ηχου1.docx")
    # Αποθήκευση του εγγράφου


def witness_domestic():
    global context
    doc = Document()
    title = doc.add_heading('ΕΚΘΕΣΗ ΕΞΕΤΑΣΗΣ ΜΑΡΤΥΡΑ ΧΩΡΙΣ ΟΡΚΟ (Κ.Π.Δ)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph = doc.add_paragraph()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run(
        '    Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}} ημέρα εβδομάδας '
        '{{day}} και ώρα {{hour}} ενώπιον εμού του {{first_officer}} του {{policeStation}} '
        'Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της ιδίας υπηρεσίας, που προσλήφθηκε '
        'ως Β\' Ανακριτικός Υπάλληλος, εμφανίστηκε ο/η κατωτέρω σημειούμενος/η μάρτυρας, ο/η οποίος/α αφού ρωτήθηκε '
        'για την ταυτότητα του/της κ.λ.π. απάντησε ότι ονομάζεται: {{surname}} {{name}} του {{fathername}} και της '
        '{{mothername}} γεν. '
        '{{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου {{tel}}, ηλεκτρονικό '
        'ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}} '
        'Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}}'
    )
    new_para1 = doc.add_paragraph(" Εξετάζεται χωρίς όρκο κατ’ εφαρμογή του άρθρου 19 παρ.1 Ν 3500/2006.")
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph("ΕΡΩΤΗΣΗ: Τι προσήλθες να καταθέσεις;")
    new_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para3 = doc.add_paragraph(
        "ΑΠΟΚΡΙΣΗ: Στις {{dateOfCrime}} και περί ώρα {{hourOfCrime}} στη {{placeOfCrime}},"
        " o/η {{stateOfVictim}}, o {{surnamePerperator}} {{namePerperator}} του "
        "{{fathernamePerperator}} και της {{mothernamePerperator}}"
        "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ."
        " {{addressPerperator}},αριθμός τηλεφώνου {{telPreperator}}, ηλεκτρονικό "
        " ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}} που εκδόθηκε την "
        "{{issuedPerperator}} από {{place_issuedPerperator}} "
        " Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}} "
        " {{whatHappened}},{{howHappened}}.{{whyHappened}}.{{add_something}}."
        "{{forensicExam}}.{{prosecution}}.ΕΡΩΤΗΣΗ:Έχουν υπάρξει ανάλογα περιστατικά στο παρελθόν;"
        "ΑΠΟΚΡΙΣΗ:{{questionPast}}. ΕΡΩΤΗΣΗ:Είναι ο/η σύζυγός σας κάτοχος κάποιου όπλου;ΑΠΟΚΡΙΣΗ:{{questionGun}}."
        "ΕΡΩΤΗΣΗ:Επιθυμείτε να ενημερώσουμε κάποια δομή παροχής αρωγής σε θύματα ενδοοικογενειακής βίας και να σας "
        "μεταφέρουμε σε κάποιο ασφαλές κατάλυμα;ΑΠΟΚΡΙΣΗ:{{safehouse}}. ΕΡΩΤΗΣΗ:Επιθυμείτε να σας μεταφέρουμε στην "
        "οικία σας;ΑΠΟΚΡΙΣΗ:{{transferHouse}}. ΕΡΩΤΗΣΗ:Επιθυμείτε να σας μεταφέρουμε στην Ιατροδικαστική Υπηρεσία;"
        "ΑΠΟΚΡΙΣΗ:{{transferForensic}}. ΕΡΩΤΗΣΗ:Έχετε κάτι άλλο να προσθέσετε;ΑΠΟΚΡΙΣΗ:{{QuestionAdd}}"
        'Ενημερώθηκα για την εφαρμογή "κομβίον πανικού".'
        "Τίποτε άλλο δεν έχω να προσθέσω και υπογράφω,")
    new_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para4 = doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hour}} ώρα και περαιώθηκε την "
                                  "{{hourOfReportFinished}} ώρα. Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού"
                                  "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5 = doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                  "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sample_document.docx')
    template = DocxTemplate("sample_document.docx")
    template.render(context)
    template.save("ΧωριςΌρκοΕνδοοικογενειακη.docx")


def witness_domestic_mutual():
    global context
    doc = Document()
    title = doc.add_heading('ΕΚΘΕΣΗ ΕΞΕΤΑΣΗΣ ΜΑΡΤΥΡΑ ΧΩΡΙΣ ΟΡΚΟ (Κ.Π.Δ)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph = doc.add_paragraph()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run(
        '    Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}} ημέρα εβδομάδας '
        '{{day}} και ώρα {{hour1}} ενώπιον εμού του {{first_officer}} του {{policeStation}} '
        'Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της ιδίας υπηρεσίας, που προσλήφθηκε '
        'ως Β\' Ανακριτικός Υπάλληλος, εμφανίστηκε ο/η κατωτέρω σημειούμενος/η μάρτυρας, ο/η οποίος/α αφού ρωτήθηκε '
        "για την ταυτότητα του/της κ.λ.π. απάντησε ότι ονομάζεται: {{surnamePerperator}} {{namePerperator}} του "
        "{{fathernamePerperator}} και της {{mothernamePerperator}}"
        "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ."
        " {{addressPerperator}},αριθμός τηλεφώνου {{telPreperator}}, ηλεκτρονικό "
        " ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}} που εκδόθηκε την "
        "{{issuedPerperator}} από {{place_issuedPerperator}} "
        " Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}} "
    )
    new_para1 = doc.add_paragraph(" Εξετάζεται χωρίς όρκο κατ’ εφαρμογή του άρθρου 19 παρ.1 Ν 3500/2006.")
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph("ΕΡΩΤΗΣΗ: Τι προσήλθες να καταθέσεις;")
    new_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para3 = doc.add_paragraph(
        "ΑΠΟΚΡΙΣΗ: Στις {{dateOfCrime1}} και περί ώρα {{hourOfCrime1}} στη {{placeOfCrime1}},"
        ' o/η {{stateOfVictim1}}, {{surname}} {{name}} του {{fathername}} και της '
        '{{mothername}} γεν. '
        '{{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου {{tel}}, ηλεκτρονικό '
        'ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}} '
        'Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}}'
        " {{whatHappened1}},{{howHappened1}}.{{whyHappened1}}.{{add_something1}}."
        "{{forensicExam1}}.{{prosecution1}}.ΕΡΩΤΗΣΗ:Έχουν υπάρξει ανάλογα περιστατικά στο παρελθόν;"
        "ΑΠΟΚΡΙΣΗ:{{questionPast1}}. ΕΡΩΤΗΣΗ:Είναι ο/η σύζυγός σας κάτοχος κάποιου όπλου;ΑΠΟΚΡΙΣΗ:{{questionGun1}}."
        "ΕΡΩΤΗΣΗ:Επιθυμείτε να ενημερώσουμε κάποια δομή παροχής αρωγής σε θύματα ενδοοικογενειακής βίας και να σας "
        "μεταφέρουμε σε κάποιο ασφαλές κατάλυμα;ΑΠΟΚΡΙΣΗ:{{safehouse1}}. ΕΡΩΤΗΣΗ:Επιθυμείτε να σας μεταφέρουμε στην "
        "οικία σας;ΑΠΟΚΡΙΣΗ:{{transferHouse1}}. ΕΡΩΤΗΣΗ:Επιθυμείτε να σας μεταφέρουμε στην Ιατροδικαστική Υπηρεσία;"
        "ΑΠΟΚΡΙΣΗ:{{transferForensic1}}. ΕΡΩΤΗΣΗ:Έχετε κάτι άλλο να προσθέσετε;ΑΠΟΚΡΙΣΗ:{{QuestionAdd1}}"
        'Ενημερώθηκα για την εφαρμογή "κομβίον πανικού".'
        "Τίποτε άλλο δεν έχω να προσθέσω και υπογράφω,")
    new_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para4 = doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hour1}} ώρα και περαιώθηκε την "
                                  "{{hourOfReportFinished1}} ώρα. Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού"
                                  "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5 = doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                  "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sample_document.docx')
    template = DocxTemplate("sample_document.docx")
    template.render(context)
    template.save("ΧωριςΌρκοΕνδοοικογενειακηmutual.docx")


def panic_button():
    global context
    # Δημιουργία εγγράφου
    doc = Document ( )

    # Δημιουργία πίνακα με 1 γραμμή και 2 κελιά (για αριστερά & δεξιά κείμενο)
    table = doc.add_table(rows = 1, cols = 2)
    table.autofit = False  # Απενεργοποίηση αυτόματου προσαρμογής πλάτους

    # Ορισμός πλάτους στηλών (50% - 50%)
    col_widths = [Inches(3), Inches(3)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Κελί 1 (Αριστερά) - "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ"
    left_cell = table.cell(0, 0)
    left_texts = [
        "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ" ,
        "ΕΛΛΗΝΙΚΗ ΑΣΤΥΝΟΜΙΑ" ,
        "ΔΙΕΥΘ. ΑΣΤ. ΘΕΣΣΑΛΟΝΙΚΗΣ" ,
        "ΑΣΤΥΝΟΜΙΚΟ ΤΜΗΜΑ ΘΕΡΜΗΣ" ,
        "ΤΗΛ:2310461203" ,
        "email: atthermis@astynomia.gr" ,
        "Αρμόδιος:{{first_officer }}" ,
        "Αρ.πρωτ: {{ protocolnumber }}"
    ]

    for text in left_texts:
        p = left_cell.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)  # γραμματοσειρά (12pt)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)

    # Κελί 2 (Δεξιά) - Τοποθεσία, ημερομηνία + επιπλέον κείμενο
    right_cell = table.cell(0, 1)

    # Γραμμή 1: Τοποθεσία & Ημερομηνία
    right_para1 = right_cell.add_paragraph()
    right_run1 = right_para1.add_run(f"Θέρμη,{datetime.now().strftime('%d/%m/%Y')}")
    right_run1.font.size = Pt(12)  # γραμματοσειρά (12pt)
    right_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT



    # --- ΚΕΝΤΡΙΚΟ ΚΕΙΜΕΝΟ (ΘΕΜΑ) ---
    doc.add_paragraph().add_run().add_break()  # Κενή γραμμή πριν το ΘΕΜΑ

    p_theme = doc.add_paragraph()
    run_theme = p_theme.add_run("ΕΝΤΥΠΟ ΕΝΗΜΕΡΩΣΗΣ ΚΑΙ ΣΥΓΚΑΤΑΘΕΣΗΣ ΧΟΡΗΓΗΣΗΣ ΤΗΣ ΕΦΑΡΜΟΓΗΣ «ΚΟΜΒΙΟΝ ΠΑΝΙΚΟΥ»"
                                "(PANIC BUTTON)")
    run_theme.bold = True
    run_theme.font.size = Pt(14)  # γραμματοσειρά (14pt)
    p_theme.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_theme.paragraph_format.space_before = Pt(16)  # Επιπλέον κενό πριν

    main_theme = doc.add_paragraph()

    offences_text1 = ", ".join(context['offences1'][:-1])
    if len(context['offences1']) > 1:
        offences_text1 += f" and {context[ 'offences1' ][ -1 ]}"
    else :
        offences_text1 = context[ 'offences1' ][ 0 ]

    run1_theme = main_theme.add_run ('O/Η κάτωθι υπογράφων/ουσα {{surname}} {{name}} του {{ fathername}} και της '
                                     '{{ mothername}} γεν. {{dateOfBirth}} στη {{ placeOfBirth }} κατ. {{ address }},'
                                     ' αριθμός τηλεφώνου {{ tel }},ηλεκτρονικό ταχυδρομείο {{ email }}, κάτοχος του '
                                     'υπ αριθμόν {{ DAT }} που εκδόθηκε την {{ issued }} από {{place_issued}} Α.Φ.Μ : '
                                     '{{ afm }}, Δ.Ο.Υ : {{ doy }}, '
                                     ' δηλώνω ότι ενημερώθηκα για τη δυνατότητα εγκατάστασης, στο κινητό'
                                     ' μου τηλέφωνο, της εφαρμογής «Κομβίον Πανικού» (Panic Button), καθώς και για τη '
                                     'διαδικασία και τους όρους εγκατάστασης και χρήσης αυτής, σύμφωνα με τα οριζόμενα '
                                     'στο άρθρο 19 του ν.4995/2022 (Α΄ 216/18-11-2022) και στην υπ’ αριθμό 111847 Κ.Υ.Α'
                                     ' (Β΄ 6007/24-11-2022), και Επιθυμώ να μου χορηγηθεί κωδικός, για πρόσβαση στην'
                                     ' εν λόγω εφαρμογή.'

         )
    run1_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    main_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν



    # Γραμμή 2: Παύλα "-Ο-" (στο κέντρο)
    p2_down = doc.add_paragraph ( )
    p2_down.add_run (
        "                                                                         \t\t-O/Η-" ).font.size = Pt (
        14 )
    p2_down.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Γραμμή 3: "Ανακριτικός Υπάλληλος" (δεξιά)
    p3_down = doc.add_paragraph ( )
    p3_down.add_run ( "\t                \t\t\tΔΙΚΑΙΟΥΧΟΣ" ).font.size = Pt ( 14 )
    p3_down.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Ορισμός γραμματοσειράς (Arial, 12pt) για όλο το κείμενο
    for row in table.rows :
        for cell in row.cells :
            for paragraph in cell.paragraphs :
                for run in paragraph.runs :
                    run.font.name = 'Arial'

        # Αποθήκευση προτύπου
        doc.save ( 'samplePb_document.docx' )
        # Επεξεργασία και τελική αποθήκευση
        template = DocxTemplate ( "samplePb_document.docx" )
        template.render ( context )
        template.save ( "panic_button.docx" )
        # Αποθήκευση του εγγράφου


def panic_button_mutual():
    global context
    # Δημιουργία εγγράφου
    doc = Document ( )

    # Δημιουργία πίνακα με 1 γραμμή και 2 κελιά (για αριστερά & δεξιά κείμενο)
    table = doc.add_table(rows = 1, cols = 2)
    table.autofit = False  # Απενεργοποίηση αυτόματου προσαρμογής πλάτους

    # Ορισμός πλάτους στηλών (50% - 50%)
    col_widths = [Inches(3), Inches(3)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Κελί 1 (Αριστερά) - "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ"
    left_cell = table.cell(0, 0)
    left_texts = [
        "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ" ,
        "ΕΛΛΗΝΙΚΗ ΑΣΤΥΝΟΜΙΑ" ,
        "ΔΙΕΥΘ. ΑΣΤ. ΘΕΣΣΑΛΟΝΙΚΗΣ" ,
        "ΑΣΤΥΝΟΜΙΚΟ ΤΜΗΜΑ ΘΕΡΜΗΣ" ,
        "ΤΗΛ:2310461203" ,
        "email: atthermis@astynomia.gr" ,
        "Αρμόδιος:{{first_officer }}" ,
        "Αρ.πρωτ: {{ protocolnumber }}"
    ]

    for text in left_texts:
        p = left_cell.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)  # γραμματοσειρά (12pt)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)

    # Κελί 2 (Δεξιά) - Τοποθεσία, ημερομηνία + επιπλέον κείμενο
    right_cell = table.cell(0, 1)

    # Γραμμή 1: Τοποθεσία & Ημερομηνία
    right_para1 = right_cell.add_paragraph()
    right_run1 = right_para1.add_run(f"Θέρμη,{datetime.now().strftime('%d/%m/%Y')}")
    right_run1.font.size = Pt(12)  # γραμματοσειρά (12pt)
    right_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT



    # --- ΚΕΝΤΡΙΚΟ ΚΕΙΜΕΝΟ (ΘΕΜΑ) ---
    doc.add_paragraph().add_run().add_break()  # Κενή γραμμή πριν το ΘΕΜΑ

    p_theme = doc.add_paragraph()
    run_theme = p_theme.add_run("ΕΝΤΥΠΟ ΕΝΗΜΕΡΩΣΗΣ ΚΑΙ ΣΥΓΚΑΤΑΘΕΣΗΣ ΧΟΡΗΓΗΣΗΣ ΤΗΣ ΕΦΑΡΜΟΓΗΣ «ΚΟΜΒΙΟΝ ΠΑΝΙΚΟΥ»"
                                "(PANIC BUTTON)")
    run_theme.bold = True
    run_theme.font.size = Pt(14)  # γραμματοσειρά (14pt)
    p_theme.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_theme.paragraph_format.space_before = Pt(16)  # Επιπλέον κενό πριν

    main_theme = doc.add_paragraph()

    offences_text1 = ", ".join(context['offences1'][:-1])
    if len(context['offences1']) > 1:
        offences_text1 += f" and {context[ 'offences1' ][ -1 ]}"
    else :
        offences_text1 = context[ 'offences1' ][ 0 ]

    run1_theme = main_theme.add_run ('O/Η κάτωθι υπογράφων/ουσα {{surnamePerperator}} {{namePerperator}} του '
                                     "{{fathernamePerperator}} και της {{mothernamePerperator}}"
                                     "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ."
                                     " {{addressPerperator}},αριθμός τηλεφώνου {{telPreperator}}, ηλεκτρονικό "
                                     " ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}} που "
                                     "εκδόθηκε την {{issuedPerperator}} από {{place_issuedPerperator}} "
                                     " Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}} "
                                     ' δηλώνω ότι ενημερώθηκα για τη δυνατότητα εγκατάστασης, στο κινητό'
                                     ' μου τηλέφωνο, της εφαρμογής «Κομβίον Πανικού» (Panic Button), καθώς και για τη '
                                     'διαδικασία και τους όρους εγκατάστασης και χρήσης αυτής, σύμφωνα με τα οριζόμενα '
                                     'στο άρθρο 19 του ν.4995/2022 (Α΄ 216/18-11-2022) και στην υπ’ αριθμό 111847 Κ.Υ.Α'
                                     ' (Β΄ 6007/24-11-2022), και Επιθυμώ να μου χορηγηθεί κωδικός, για πρόσβαση στην'
                                     ' εν λόγω εφαρμογή.'

         )
    run1_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    main_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν



    # Γραμμή 2: Παύλα "-Ο-" (στο κέντρο)
    p2_down = doc.add_paragraph ( )
    p2_down.add_run (
        "                                                                         \t\t-O/Η-" ).font.size = Pt (
        14 )
    p2_down.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Γραμμή 3: "Ανακριτικός Υπάλληλος" (δεξιά)
    p3_down = doc.add_paragraph ( )
    p3_down.add_run ( "\t                \t\t\tΔΙΚΑΙΟΥΧΟΣ" ).font.size = Pt ( 14 )
    p3_down.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Ορισμός γραμματοσειράς (Arial, 12pt) για όλο το κείμενο
    for row in table.rows :
        for cell in row.cells :
            for paragraph in cell.paragraphs :
                for run in paragraph.runs :
                    run.font.name = 'Arial'

        # Αποθήκευση προτύπου
        doc.save ( 'samplePbmutual_document.docx' )
        # Επεξεργασία και τελική αποθήκευση
        template = DocxTemplate ( "samplePbmutual_document.docx" )
        template.render ( context )
        template.save ( "panic_button_mutual.docx" )
        # Αποθήκευση του εγγράφου



                                  