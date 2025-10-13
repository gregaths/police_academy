# Note: This is the complete Python code for the requested application, structured into separate files as specified.
# Each button/class is in its own file. The documents generation is in a separate module.
# Assumptions:
# - Images (image1.png, image2.png) are local files in the working directory. image1.png is set as background or label in main window if needed, but since not viewable, used text buttons.
# - Bio is assumed to be a simple text window; if image2.png is the bio, it loads as QLabel.
# - Fixed typos in variable names (removed spaces in keys like 'dateOfCrime ', 'telPreperator ').
# - Used QLineEdit for all inputs, QTextEdit for longer fields like whatHappened.
# - For offences, used QListWidget with add button.
# - Only implementing case 2a as specified: known perpetrator, arrested (no mutual for now).
# - User selects known and arrested via radio buttons.
# - Global context dictionary.
# - Requires pip install python-docx docxtpl pyqt5 (not in tool, but for user).
# - Greek characters supported with utf-8.

# documents.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docxtpl import DocxTemplate
from datetime import datetime

context = {}


def witness_report():
    global context
    doc = Document()
    title = doc.add_heading('ΕΚΘΕΣΗ ΕΝΟΡΚΗΣ ΕΞΕΤΑΣΗΣ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph = doc.add_paragraph()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run(
        '    Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}} ημέρα εβδομάδας '
        '{{day}} και ώρα {{hour}} ενώπιον εμού του {{first_officer}} του {{policeStation}} '
        'Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της ιδίας υπηρεσίας, που προσλήφθηκε '
        'ως Β\' Ανακριτικός Υπάλληλος, εμφανίστηκε ο κατωτέρω μάρτυρας, ο οποίος αφού ρωτήθηκε για την ταυτότητα '
        'του κ.λ.π. απάντησε ότι ονομάζεται: {{surname}} {{name}} του {{fathername}} και της {{mothername}} γεν. '
        '{{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου {{tel}}, ηλεκτρονικό '
        'ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}} '
        'Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}}'
    )
    new_para1 = doc.add_paragraph("   Έπειτα ο εξεταζόμενος έδωσε τον προβλεπόμενο από τα άρθρα 219 και 220 παρ. 1 "
                                  " του Κώδικα Ποινικής Δικονομίας όρκο, ως ακολούθως: «Δηλώνω, επικαλούμενος την "
                                  "τιμή και την συνείδηση μου, ότι θα πω όλη την αλήθεια και μόνο την αλήθεια, χωρίς"
                                  " να προσθέσω ούτε να αποκρύψω τίποτα», και στην συνέχεια εξετάσθηκε ως εξής:.")
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph("ΕΡΩΤΗΣΗ: Ρωτήθηκε σχετικά:")
    new_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para3 = doc.add_paragraph(
        "ΑΠΟΚΡΙΣΗ: Στις {{dateOfCrime}} και περί ώρα {{hourOfCrime}} στη {{placeOfCrime}},"
        " {{stateOfVictim}}, o {{surnamePerperator}} {{namePerperator}} του "
        "{{fathernamePerperator}} και της {{mothernamePerperator}}"
        "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ."
        " {{addressPerperator}},αριθμός τηλεφώνου {{telPreperator}}, ηλεκτρονικό "
        " ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}} που εκδόθηκε την "
        "{{issuedPerperator}} από {{place_issuedPerperator}} "
        " Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}} "
        " {{whatHappened}},{{howHappened}}.{{whyHappened}}.{{add_something}}. {{forensicExam}}.{{prosecution}}."
        "Τίποτε άλλο δεν έχω να προσθέσω και υπογράφω,")
    new_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para4 = doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hour}} ώρα και περαιώθηκε την "
                                  "{{hourOfReportFinished}} ώρα. Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού"
                                  "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5 = doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                  "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sample_document.docx')
    template = DocxTemplate("sample_document.docx")
    template.render(context)
    template.save("output1.docx")

def iatrodik():
    global context
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    col_widths = [Inches(3), Inches(3)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
    left_cell = table.cell(0, 0)
    left_texts = [
        "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ",
        "ΕΛΛΗΝΙΚΗ ΑΣΤΥΝΟΜΙΑ",
        "ΔΙΕΥΘ. ΑΣΤ. ΘΕΣΣΑΛΟΝΙΚΗΣ",
        "ΑΣΤΥΝΟΜΙΚΟ ΤΜΗΜΑ ΘΕΡΜΗΣ",
        "ΤΗΛ:2310461203",
        "email: atthermis@astynomia.gr",
        "Αρμόδιος:{{first_officer}}",
        "Αρ.πρωτ: {{protocolnumber}}"
    ]
    for text in left_texts:
        p = left_cell.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
    right_cell = table.cell(0, 1)
    right_para1 = right_cell.add_paragraph()
    right_run1 = right_para1.add_run(f"Θέρμη, {datetime.now().strftime('%d/%m/%Y')}")
    right_run1.font.size = Pt(12)
    right_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para2 = right_cell.add_paragraph()
    right_run2 = right_para2.add_run(" Π Ρ Ο Σ")
    right_run2.font.size = Pt(12)
    right_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_para3 = right_cell.add_paragraph()
    right_run3 = right_para3.add_run(
        "ΕΡΓΑΣΤΗΡΙΟ ΙΑΤΡΟΔΙΚΑΣΤΙΚΗΣ ΚΑΙ\nΤΟΞΙΚΟΛΟΓΙΑΣ\nΑριστοτέλειο Πανεπιστήμιο Θεσσαλονίκης")
    right_run3.font.size = Pt(12)
    right_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run().add_break()
    p_theme = doc.add_paragraph()
    run_theme = p_theme.add_run("ΘΕΜΑ: «Παραγγελία Ιατροδικαστικής εξέτασης»")
    run_theme.bold = True
    run_theme.font.size = Pt(14)
    p_theme.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_theme.paragraph_format.space_before = Pt(16)
    main_theme = doc.add_paragraph()
    offences_text = ", ".join(context['offences'][:-1])
    if len(context['offences']) > 1:
        offences_text += f" και {context['offences'][-1]}"
    else:
        offences_text = context['offences'][0] if context['offences'] else ""
    context['offences_text'] = offences_text  # Add to context if needed, but code uses it directly
    run1_theme = main_theme.add_run(
        "    Παρακαλούμε μεριμνήσετε για την εξέταση από Ιατροδικαστή της υπηρεσίας σας, "
        'του {{surname}} {{name}} του {{fathername}} και της {{mothername}} γεν. '
        '{{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου '
        '{{tel}}, ηλεκτρονικό ταχυδρομείο {{email}},Α.Φ.Μ :{{afm}}, Δ.Ο.Υ :{{doy}} '
        'διότι την {{dateOfCrime}}  και ώρα {{hourOfCrime}}επί της {{placeOfCrime}} '
        'δέχτηκε σωματικές βλάβες από τον '
        "{{surnamePerperator}} {{namePerperator}} του {{fathernamePerperator}} και της "
        " {{mothernamePerperator}} γεν. {{dateOfBirthPerperator}} στη "
        "{{placeOfBirthPerperator}} κατ.{{addressPerperator}},αριθμός τηλεφώνου "
        "{{telPreperator}}, ηλεκτρονικό  ταχυδρομείου{{emailPreperator}}, Α.Φ.Μ : "
        "{{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}}όπως ρητά δηλώθηκε στην από "
        " {{date_num}} {{month}}  {{year}} και "
        "ώρα {{hour}} έγκληση σε βάρος του προρριθέντα κατά παράβαση του  άρθρου "
        "{{offences_text}} του Π.Κ. Συγκεκριμένα "
        "και όπως μας δήλωσε παθόντας προκλήθηκαν σωματικές βλάβες  {{placeOfHurt}}.")
    run1_theme.font.size = Pt(14)
    main_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_theme.paragraph_format.space_before = Pt(16)
    main2_theme = doc.add_paragraph()
    run2_theme = main2_theme.add_run(
        "Παρακαλούμε όπως με σχετική έκθεση σας μας απαντήσετε στα κατωτέρω ερωτήματα μας.")
    run2_theme.font.size = Pt(14)
    main2_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main2_theme.paragraph_format.space_before = Pt(16)
    list_items = [
        "Φέρει σωματικές  κακώσεις και ποιες",
        "Προ πόσου χρόνου προξενήθηκαν αυτές.",
        "Με ποιο όργανο προξενήθηκαν.",
        "Η σωματική κάκωση ή βλάβη είναι απλή.",
        "Η σωματική βλάβη είναι επικίνδυνη δηλαδή τελέστηκε με τρόπο που θα μπορούσε να προκαλέσει κίνδυνο ζωής του "
        "παθόντος ή βαριά σωματική βλάβη.",
        "Ότι άλλο  από  την επιστήμη σας δύνασθε να μας γνωρίσετε. "
    ]
    for item in list_items:
        p_list = doc.add_paragraph(style='List Number')
        p_list.add_run(item).font.size = Pt(14)
        p_list.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main3_theme = doc.add_paragraph()
    run3_theme = main3_theme.add_run(
        "Την έκθεση πραγματογνωμοσύνης να μας την εγχειρίσετε το συντομότερο δυνατόν.")
    run3_theme.font.size = Pt(14)
    main3_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main3_theme.paragraph_format.space_before = Pt(16)
    doc.add_paragraph().add_run().add_break()
    p1_down = doc.add_paragraph()
    p1_down.add_run(
        "ΕΚΤΕΛΕΣΘΗΚΕ ΚΑΙ ΚΑΤΑΤΕΘΗΚΕ ΠΙΝΑΚΑΣ\nΙΑΤΡΟΔΙΚΑΣΤΙΚΩΝ ΔΙΚΑΙΩΜΑΤΩΝ\n\tΑπό Ευρώ.............................."
        "..........").font.size = Pt(14)
    p1_down.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2_down = doc.add_paragraph()
    p2_down.add_run(
        "                              -Ο-                                         \t\t-O-").font.size = Pt(
        14)
    p2_down.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3_down = doc.add_paragraph()
    p3_down.add_run("\tΑνακριτικός Υπάλληλος\t\t\tΑνακριτικός Υπάλληλος").font.size = Pt(14)
    p3_down.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
    doc.save('sampleiatr_document.docx')
    template = DocxTemplate("sampleiatr_document.docx")
    template.render(context)
    template.save("iatrodikastikh.docx")

# Similar for other functions: silipsi, deltio_kat, apologia, rights, vas, transmission
# (To save space, I'm abbreviating the remaining functions; they are copied from the user query with fixed keys)


def silipsi():
    global context
    doc = Document()
    title = doc.add_heading('ΕΚΘΕΣΗ ΣΥΛΛΗΨΗΣ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph = doc.add_paragraph()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    offences_text = ", ".join(context['offences'][:-1])
    if len(context['offences']) > 1:
        offences_text += f" and {context['offences'][-1]}"
    else:
        offences_text = context['offences'][0] if context['offences'] else ""
    main_paragraph.style.font.size = Pt(14)
    main_paragraph.paragraph_format.space_before = Pt(16)
    main_paragraph.add_run(
        '    Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}} ημέρα εβδομάδας '
        '{{day}} και ώρα {{startTimeArrestReport}} ενώπιον εμού του {{first_officer}} του {{policeStation}} '
        'Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της ιδίας υπηρεσίας, που προσλήφθηκε '
        'ως Β\' Ανακριτικός Υπάλληλος, οδηγήθηκε στο κατάστημα του Α.Τ. Θέρμης, ο κάτωθι {{surnamePerperator}} '
        ' {{namePerperator}} του {{fathernamePerperator}} και της {{mothernamePerperator}}'
        ' γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ.'
        " {{addressPerperator}},αριθμός τηλεφώνου {{telPreperator}}, ηλεκτρονικό "
        " ταχυδρομείου{{emailPreperator}}, κάτοχος του υπ αριθμόν {{DATperperator}} που εκδόθηκε την "
        "{{issuedPerperator}} από {{place_issuedPerperator}}Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}}, "
        "από τον {{officer_arrest}} υπηρετών στο {{policeStation}}, που τον συνέλαβε στις {{dateOfArrest}} "
        "και ώρα {{hourOfArrest}} στη {{placeOfArrest}} για παράβαση του/των άρθρων {{offences_text}} του Π.Κ."
    )
    main2_theme = doc.add_paragraph()
    run2_theme = main2_theme.add_run(
        "  Αφού εξέτασα αυτήν, διέταξα την παραπομπή και παράδοση της με την παρούσα έκθεση στον  κ. Εισαγγελέα"
        " Πλημμελειοδικών Θεσσαλονίκης.")
    run2_theme.font.size = Pt(14)
    main2_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main2_theme.paragraph_format.space_before = Pt(16)
    main3_theme = doc.add_paragraph()
    run3_theme = main3_theme.add_run('Η παρούσα έκθεση άρχισε να συντάσσεται την {{startTimeArrestReport}} ’ '
                                     'ώρα και περατώθηκε την {{endTimeOfReport}} ώρα.  '
                                     )
    run3_theme.font.size = Pt(14)
    main3_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main3_theme.paragraph_format.space_before = Pt(16)
    new_para4 = doc.add_paragraph(" Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού"
                                  "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5 = doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                  "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sampleArrest_document.docx')
    template = DocxTemplate("sampleArrest_document.docx")
    template.render(context)
    template.save("ΕΚΘΕΣΗ_ΣΥΛΛΗΨΗΣ.docx")


def deltio_kat():
    global context

    # Δημιουργία εγγράφου
    doc = Document ( )
    title = doc.add_heading ( 'ΔΕΛΤΙΟ ΣΤΟΙΧΕΙΩΝ ΤΑΥΤΟΤΗΤΑΣ ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ' , level = 0 )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ορισμός γραμματοσειράς (bold + υπογραμμισμένο) για τον τίτλο
    for run in title.runs:
        run.font.size = Pt ( 14 )
        run.bold = True  # Bold
        run.underline = WD_UNDERLINE.SINGLE  # Υπογράμμιση

        # ΕΠΩΝΥΜΟ
        p_eponymo = doc.add_paragraph ( )
        run_label = p_eponymo.add_run ( "ΕΠΩΝΥΜΟ: " )
        run_label.bold = True
        run_value = p_eponymo.add_run ( " {{ surnamePerperator }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝΟΜΑ
        p_onoma = doc.add_paragraph ( )
        run_label = p_onoma.add_run ( "ΟΝΟΜΑ:    " )
        run_label.bold = True
        run_value = p_onoma.add_run ( "{{ namePerperator  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝΟΜΑ ΠΑΤΕΡΑ
        p_onoma_patera = doc.add_paragraph ( )
        run_label = p_onoma_patera.add_run ( "ΟΝΟΜΑ ΠΑΤΕΡΑ:    " )
        run_label.bold = True
        run_value = p_onoma_patera.add_run ( "{{ fathernamePerperator }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝΟΜΑ ΜΗΤΕΡΑΣ
        p_onoma_mhteras = doc.add_paragraph ( )
        run_label = p_onoma_mhteras.add_run ( "ΟΝΟΜΑ ΜΗΤΕΡΑΣ:    " )
        run_label.bold = True
        run_value = p_onoma_mhteras.add_run ( " {{ mothernamePerperator }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΠΑΤΡΙΚΟ ΕΠΩΝΥΜΟ (Για έγγαμες)
        p_onoma_patriko = doc.add_paragraph ( )
        run_label = p_onoma_patriko.add_run ( "ΠΑΤΡΙΚΟ ΕΠΩΝΥΜΟ (Για έγγαμεσ):    " )
        run_label.bold = True
        run_value = p_onoma_patriko.add_run ( "_______________________" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝ/ΜΟ ΣΥΖΥΓΟΥ
        p_onoma_syzygou = doc.add_paragraph ( )
        run_label = p_onoma_syzygou.add_run ( "ΟΝ/ΜΟ ΣΥΖΥΓΟΥ:    " )
        run_label.bold = True
        run_value = p_onoma_syzygou.add_run ( "_______________________" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΗΜΕΡΟΜΗΝΙΑ ΓΕΝΝΗΣΗΣ
        p_dateofbirth = doc.add_paragraph ( )
        run_label = p_dateofbirth.add_run ( "ΗΜΕΡΟΜΗΝΙΑ ΓΕΝΝΗΣΗΣ:    " )
        run_label.bold = True
        run_value = p_dateofbirth.add_run ( "  {{ dateOfBirthPerperator   }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΤΟΠΟΘΕΣΙΑ ΓΕΝΝΗΣΗΣ
        p_placeofbirth = doc.add_paragraph ( )
        run_label = p_placeofbirth.add_run ( " ΤΟΠΟΘΕΣΙΑ ΓΕΝΝΗΣΗΣ:    " )
        run_label.bold = True
        run_value = p_placeofbirth.add_run ( "  {{ placeOfBirthPerperator   }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΤΟΠΟΘΕΣΙΑ κατοικιασ
        p_address = doc.add_paragraph ( )
        run_label = p_address.add_run ( " ΤΟΠΟΣ ΚΑΤΟΙΚΙΑΣ:    " )
        run_label.bold = True
        run_value = p_address.add_run ( "  {{ addressPerperator  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # Δ.Α.Τ
        p_dat = doc.add_paragraph ( )
        run_label = p_dat.add_run ( " ΑΡΙΘΜΟΣ Δ.Α.Τ:    " )
        run_label.bold = True
        run_value = p_dat.add_run ( "  {{ DATperperator  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # Δ.Α.Τ
        p_issue = doc.add_paragraph ( )
        run_label = p_issue.add_run ( " ΕΚΔΟΘΕΝ:    " )
        run_label.bold = True
        run_value = p_issue.add_run ( "  {{ issuedPerperator  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # Δ.Α.Τ
        p_issueplace = doc.add_paragraph ( )
        run_label = p_issueplace.add_run ( " ΑΠΟ:    " )
        run_label.bold = True
        run_value = p_issueplace.add_run ( " {{ place_issuedPerperator  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΑΦΜ
        # Δ.Α.Τ
        p_afm = doc.add_paragraph ( )
        run_label = p_afm.add_run ( " Α.Φ.Μ:    " )
        run_label.bold = True
        run_value = p_afm.add_run ( "  {{ afmPreperator }} από Δ.Ο.Υ {{  doyPrep }}  " )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        footer_para = doc.add_paragraph ( )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Προσθήκη κειμένου με ακριβή θέση (χρήση tabs)
        footer_run = footer_para.add_run ( )
        footer_run.add_text ( f"\t\tΘέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
        footer_run.add_text ( "\t\t\tΟ\t\t\n" )  # 5 tabs για στοίχιση
        footer_run.add_text ( "\t\t\t\t\tΑνακριτικός Υπάλληλος\n\n" )
        footer_run.add_text ( "\t\t\t\t\t {{ first_officer }}\n" )

        # Αποθήκευση
        doc.save ( 'deltio_katigoroumenoutemplate.docx' )
        # Επεξεργασία και τελική αποθήκευση
        template = DocxTemplate ( 'deltio_katigoroumenoutemplate.docx' )
        template.render ( context )
        template.save ( "ΔΕΛΤΙΟ_ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ.docx" )
        # Αποθήκευση του εγγράφου


def apologia():
        global context  # Declare we're using the global context
        # Δημιουργία εγγράφου
        doc = Document()
        # Προσθήκη τίτλου (κεντραρισμένου)
        title = doc.add_heading('ΕΚΘΕΣΗ ΕΞΕΤΑΣΗΣ ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
        main_paragraph = doc.add_paragraph()
        main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph.add_run('Στην {{ place }} σήμερα την {{ date_num_apologia}} του μήνα {{ month_apologia }} του'
                               ' έτους {{year_apologia}} ημέρα εβδομάδας {{ day_apologia}} και ώρα {{start_hour_apologia}}'
                               ' ενώπιον εμού του {{ first_officer }} του {{ policeStation }} Θεσσαλονίκης, παρισταμένου  '
                               'και του  {{ sec_officer }} της ιδίας υπηρεσίας, που προσλήφθηκε ως Β\' Ανακριτικός'
                               ' Υπάλληλος, εξετάζεται ο κατωτέρω σημειούμενος κατηγορούμενος ,')

        # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
        main_paragraph1 = doc.add_paragraph()
        main_paragraph1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph1.add_run('ΕΡΩΤΗΣΗ : Πως ονομάζεσαι κ.λ.π.')

        # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
        main_paragraph2 = doc.add_paragraph()
        main_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph2.add_run(
            'ΑΠΟΚΡΙΣΗ : {{surnamePerperator}}  {{namePerperator}} του {{fathernamePerperator}} και της '
            '{{ mothernamePerperator}} γεν. {{dateOfBirthPerperator }} στη {{ placeOfBirthPerperator }}'
            ' κατ.{{ addressPerperator }},αριθμός τηλεφώνου {{ telPreperator }}, ηλεκτρονικό '
            'ταχυδρομείου{{ emailPreperator }}, κάτοχος του υπ αριθμόν {{ DATperperator }} που '
            'εκδόθηκε την {{ issuedPerperator }} από {{place_issuedPerperator}}Α.Φ.Μ : '
            '{{ afmPreperator }}, Δ.Ο.Υ : {{ doyPrep }}, από τον {{officer_arrest }} υπηρετών στο '
            ' {{  policeStation }}')

        # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
        main_paragraph3 = doc.add_paragraph()
        main_paragraph3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph3.add_run('Ενταύθα γνωρίσαμε στον εξεταζόμενο ότι κατηγορείται για παράβαση του άρθρου '
                                '{{offences }} του Π.Κ.')

        # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
        main_paragraph4 = doc.add_paragraph()
        main_paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph4.add_run('Ειδικότερα κατηγορείται ότι στις {{dateOfCrime}}και περί ώρα {{ hourOfCrime}}'
                                ' στο {{placeOfCrime}}  προέβης στην παράβαση του/των άρθρων {{offences }} τπυ Π.Κ.'
                                'εις βάρος του {{surname}} {{name}} του {{ fathername}} και της {{ mothername}} γεν.'
                                ' {{dateOfBirth  }} στη {{ placeOfBirth }} κατ. {{ address }}, αριθμός τηλεφώνου '
                                '{{ tel }}, ηλεκτρονικό ταχυδρομείο {{ email }}, κάτοχος του υπ αριθμόν {{ DAT }} που '
                                'εκδόθηκε την  {{ issued }} από {{place_issued}} Α.Φ.Μ : {{ afm }}, Δ.Ο.Υ : {{ doy }}')

        # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
        main_paragraph5 = doc.add_paragraph()
        main_paragraph5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph5.add_run(
            'Στη συνέχεια αφού ανακοινώσαμε στον εξεταζόμενο το περιεχόμενο των εγγράφων της ανάκρισης,'
            ' εξηγήσαμε εκ νέου με σαφήνεια σ’ αυτόν , σύμφωνα με τα άρθρα 95 και 105 ή 106 του Κώδικα '
            'Ποινικής Δικονομίας , όλα τα δικαιώματά του που προβλέπονται από τα άρθρα  91,95,96,97,98,'
            '100,101 , 104 και 273 του Κώδικα Ποινικής Δικονομίας και ειδικότερα ότι δικαιούται : να '
            'παρίσταται μετά συνηγόρου , μετά του οποίου δεν μπορεί να απαγορευτεί η απολύτως απόρρητη '
            'επικοινωνία του σε καμιά περίπτωση, να μελετήσει ο ίδιος ή ο συνήγορός του  τα έγγραφα της'
            ' ανάκρισης και του κατηγορητηρίου (πραγματικών περιστατικών- κατηγορίας) , να του '
            'χορηγηθούν αντίγραφα αυτών με δική του δαπάνη και μετά από γραπτή αίτησή του να ζητήσει '
            'προθεσμία μέχρι 48 ωρών , προ της παρέλευσης της οποίας δεν υποχρεούται σε απολογία και '
            'ότι δύναται να δοθεί παράταση της προθεσμίας αυτής με αίτησή του , να αρνηθεί να απαντήσει'
            ' (δικαίωμα σιωπής και μη αυτοενεχοποίσης) , και να παραδώσει την απολογία του γραπτή . ')

        main_paragraph6 = doc.add_paragraph()
        main_paragraph6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph6.add_run('ΕΡΩΤΗΣΗ: Επιθυμείται να κάνετε χρήση των δικαιωμάτων που σας γνωστοποιήθηκαν ;  ')

        main_paragraph6 = doc.add_paragraph()
        main_paragraph6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph6.add_run('ΑΠΟΚΡΙΣΗ: {{ question_rights_apologia }}  ')

        main_paragraph7 = doc.add_paragraph()
        main_paragraph7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph7.add_run('ΑΠΟΚΡΙΣΗ: {{ question_rights_apologia }}  ')

        main_paragraph8 = doc.add_paragraph()
        main_paragraph8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph8.add_run('Ύστερα από τα ανωτέρω προβήκαμε στην λήψη της απολογίας του  ')

        main_paragraph9 = doc.add_paragraph()
        main_paragraph9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph9.add_run('Για πιστοποίηση συντάχθηκε  η παρούσα έκθεση, η οποία αφού αναγνώσθηκε και βεβαιώθηκε '
                                ' υπογράφεται ως ακολούθως : ')

        main_paragraph10 = doc.add_paragraph()
        main_paragraph10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph10.add_run('Ο Εξετασθείς                              Ο Β Ανακριτικός Υπάλληλος '
                                 '              Ο   Ανακριτικός Υπάλληλος\n\n\n')

        main_paragraph11 = doc.add_paragraph()
        main_paragraph11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph11.add_run(' Στη συνέχεια, προβήκαμε στην εξέταση του κατηγορουμένου ως ακολούθως : ')

        main_paragraph12 = doc.add_paragraph()
        main_paragraph12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph12.add_run(' ΕΡΩΤΗΣΗ :  Κατηγορήθηκες  άλλη φορά και για ποια αιτία ;')

        main_paragraph13 = doc.add_paragraph()
        main_paragraph13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph13.add_run('ΑΠΟΚΡΙΣΗ: {{ question_crime_apologia }}  ')

        main_paragraph14 = doc.add_paragraph()
        main_paragraph14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph14.add_run('ΕΡΩΤΗΣΗ : Κατηγορείσαι ήδη για τις πράξεις που σου γνωστοποιήθηκαν ανωτέρω. '
                                 'Τι απολογείσαι;   ')

        main_paragraph15 = doc.add_paragraph()
        main_paragraph15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
        main_paragraph15.add_run('ΑΠΟΚΡΙΣΗ : {{question_crime_confess }} ')

        new_para4 = doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την  {{start_hour_apologia}} ώρα και "
                                      "περαιώθηκε την {{end_hour_apologia}} ώρα. Για πίστωση συντάχθηκε η παρούσα "
                                      "έκθεση η οποία αφού αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
        new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        new_para5 = doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                      "         Ο   Ανακριτικός Υπάλληλος")
        new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.save('sample_confess.docx')
        # Επεξεργασία και τελική αποθήκευση
        template = DocxTemplate("sample_confess.docx")
        template.render(context)
        template.save("ΕΚΘΕΣΗ_ΕΞΕΤΑΣΗΣ_ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ.docx")
        # Αποθήκευση του εγγράφου


def rights():
    global context  # Declare we're using the global context
    # Δημιουργία εγγράφου
    doc = Document ()
    # Προσθήκη τίτλου (κεντραρισμένου)
    title = doc.add_heading ( 'ΕΚΘΕΣΗ ΕΞΕΤΑΣΗΣ ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ' , level = 0 )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph = doc.add_paragraph ()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph.add_run ( 'Στην {{ place }} σήμερα την {{ date_num_apologia}} του μήνα {{ month_apologia }} του'
                             ' έτους {{year_apologia}} ημέρα εβδομάδας {{ day_apologia}} και ώρα {{start_time_rights}}'
                             ' ενώπιον εμού του {{ first_officer }} του {{ policeStation }} Θεσσαλονίκης, παρισταμένου  '
                             'και του  {{ sec_officer }} της ιδίας υπηρεσίας, που προσλήφθηκε ως Β\' Ανακριτικός'
                             ' Υπάλληλος, προσκλήθηκε ο  κατωτέρω σημειούμενος  κατηγορούμενος, που ονομάζεται '
                             ' {{surnamePerperator}} {{namePerperator}} του {{fathernamePerperator}} και της '
                             '{{ mothernamePerperator}} γεν. {{dateOfBirthPerperator }} στη {{ placeOfBirthPerperator }}'
                             ' κατ.{{ addressPerperator }},αριθμός τηλεφώνου {{ telPreperator }}, ηλεκτρονικό  '
                             'ταχυδρομείου{{ emailPreperator }}, κάτοχος του υπ αριθμόν {{ DATperperator }} που '
                             'εκδόθηκε την {{ issuedPerperator }} από {{place_issuedPerperator}}Α.Φ.Μ : '
                             '{{ afmPreperator }}, Δ.Ο.Υ : {{ doyPrep }},στον οποίο γνωστοποιήσαμε ότι κατηγορείται '
                             'για παράβαση του/των άρθρων  {{offences }} τπυ Π.Κ. και εξηγήσαμε με σαφήνεια και'
                             ' πληρότητα σ’ αυτόν βάσει του άρθρου 95 του Κώδικα Ποινικής  Δικονομίας όλα τα εκ των '
                             ' άρθρων 91,95,96,97,98,99,100,103 και 104  του  Κ.Π.Δ.  δικαιώματά  του/της  και  '
                             'αναλυτικότερα :' )

    main_paragraph1 = doc.add_paragraph ()
    main_paragraph1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph1.add_run ( 'Α. Εκ του άρθρου 91 του Κ.Π.Δ. , ήτοι δικαίωμα παροχής δωρεάν νομικής βοήθεια :' )

    main_paragraph2 = doc.add_paragraph ()
    main_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph2.add_run (
        'Ο ύποπτος ή ο κατηγορούμενος έχουν δικαίωμα δωρεάν νομικής βοήθειας, το οποίο περιλαμβάνει'
        ' παροχή νομικών συμβουλών και νομική αρωγή και εκπροσώπηση τους ενώπιον του δικαστηρίου , '
        'σύμφωνα με όσα ορίζονται σε σχετικές διατάξεις. ' )

    main_paragraph2 = doc.add_paragraph ()
    main_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph2.add_run ( 'B. Εκ του άρθρου 95 του Κ.Π.Δ. , ήτοι δικαίωμα σε ενημέρωση :' )

    main_paragraph3 = doc.add_paragraph ()
    main_paragraph3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph3.add_run (
        '1.	Ο ύποπτος ή ο κατηγορούμενος ενημερώνεται αμέσως όσον αφορά τουλάχιστον στα ακόλουθα '
        'δικαιώματα: α) το δικαίωμα παράστασης με συνήγορο, β) το δικαίωμα και τις προϋποθέσεις '
        'παροχής δωρεάν νομικών συμβουλών, γ) το δικαίωμα ενημέρωσης σχετικά με την κατηγορία, '
        'δ) το δικαίωμα διερμηνείας και μετάφρασης και ε) το δικαίωμα σιωπής και μη '
        'αυτοενοχοποίησης. 2. Η ενημέρωση σύμφωνα με την παρ. 1 παρέχεται σε απλή και κατανοητή '
        'γλώσσα, προφορικώς ή εγγράφως, λαμβανομένων υπόψη των ειδικών αναγκών των υπόπτων ή '
        'κατηγορουμένων που είναι ευάλωτα πρόσωπα. Αντικείμενο της ενημέρωσης οφείλει να αποτελεί '
        'και η αναφορά των συνεπειών παραίτησης από την άσκηση των δικαιωμάτων. Για την ενημέρωση '
        'και την απάντηση του υπόπτου ή του κατηγορουμένου συντάσσεται έκθεση, η οποία και '
        'υπογράφεται.' )

    main_paragraph4 = doc.add_paragraph ()
    main_paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph4.add_run ( 'Γ. Εκ του άρθρου 96 , ήτοι δικαίωμα χορήγησης εγγράφου περί των δικαιωμάτων : ' )

    main_paragraph5 = doc.add_paragraph ()
    main_paragraph5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph5.add_run (
        '1.	Στον ύποπτο ή τον κατηγορούμενο, ο οποίος συλλαμβάνεται ή κρατείται, παρέχεται αμέσως '
        'έγγραφο στο οποίο καταγράφονται τα δικαιώματά του και του επιτρέπεται να το διατηρεί στην'
        ' κατοχή του καθ’ όλη τη διάρκεια της στέρησης της ελευθερίας του. Το έγγραφο αυτό '
        'περιλαμβάνει πληροφορίες σχετικά με τα ακόλουθα δικαιώματα: α) το δικαίωμα παράστασης με '
        'συνήγορο, β) το δικαίωμα και τις προϋποθέσεις παροχής δωρεάν νομικών συμβουλών, γ) το '
        'δικαίωμα ενημέρωσης σχετικά με την κατηγορία, δ) το δικαίωμα διερμηνείας και μετάφρασης, '
        'ε) το δικαίωμα σιωπής και μη αυτοενοχοποίησης, στ) το δικαίωμα πρόσβασης στο υλικό της'
        ' δικογραφίας, ζ) το δικαίωμα ενημέρωσης των προξενικών αρχών και ενός επιπλέον προσώπου '
        'της επιλογής του, η) το δικαίωμα σε επείγουσα ιατρική περίθαλψη, θ) τον ανώτατο αριθμό'
        ' ωρών ή ημερών κατά τις οποίες ο κατηγορούμενος δύναται να στερηθεί της ελευθερίας του '
        'προτού προσαχθεί ενώπιον δικαστικής αρχής και ι) πληροφορίες σχετικά με τις δυνατότητες '
        'προσβολής του νόμιμου χαρακτήρα της σύλληψης ή της κράτησης. 2. Το ως άνω έγγραφο '
        'συντάσσεται σε απλή και κατανοητή γλώσσα. Όταν αυτό δεν είναι διαθέσιμο στην κατάλληλη '
        'γλώσσα, ο ύποπτος ή ο κατηγορούμενος ενημερώνεται για τα δικαιώματά του προφορικά σε '
        'γλώσσα που κατανοεί. Το εν λόγω έγγραφο πρέπει στη συνέχεια να χορηγείται, χωρίς'
        ' αδικαιολόγητη καθυστέρηση, σε γλώσσα που ο ύποπτος ή ο κατηγορούμενος κατανοεί.' )

    main_paragraph5 = doc.add_paragraph ()
    main_paragraph5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph5.add_run (
        'Δ. Εκ του άρθρου 97 , ήτοι δικαιώμα ενημέρωσης προσώπου της επιλογής του κατηγορουμένου '
        'σε περίπτωση στέρησης της ελευθερίας του : ' )

    main_paragraph6 = doc.add_paragraph ()
    main_paragraph6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph6.add_run ( '1. Ο κατηγορούμενος έχει δικαίωμα να ζητήσει να ενημερωθεί, χωρίς αδικαιολόγητη'
                              ' καθυστέρηση, για τη στέρηση της ελευθερίας του ένα τουλάχιστον πρόσωπο της επιλογής του. '
                              'Αν ο κατηγορούμενος είναι ανήλικος ενημερώνεται ο ασκών τη γονική μέριμνα, εκτός αν αυτό '
                              'αντιβαίνει στο συμφέρον του ανηλίκου, οπότε ενημερώνεται κάποιο άλλο ενδεδειγμένο ενήλικο '
                              'πρόσωπο ή η αρμόδια για την προστασία των ανηλίκων αρχή. 2. Σε εξαιρετικές περιπτώσεις και'
                              'για να αποτραπεί άμεσος κίνδυνος είτε για τη ζωή, την ελευθερία ή τη σωματική ακεραιότητα '
                              'κάποιου προσώπου είτε για την αποδεικτική διερεύνηση του εγκλήματος, οι αρμόδιες αρχές '
                              'μπορούν να μην επιτρέπουν προσωρινά την ενημέρωση τρίτου προσώπου για τη στέρηση της '
                              'ελευθερίας του κατηγορουμένου. Στην περίπτωση αυτή εξετάζεται αν ένα άλλο τρίτο πρόσωπο, '
                              'που έχει υποδειχθεί από τον κατηγορούμενο, μπορεί να ενημερωθεί σχετικά. Αν ο'
                              ' κατηγορούμενος είναι ανήλικος, η αρμόδια για την προστασία των ανηλίκων αρχή ενημερώνεται'
                              'στην περίπτωση αυτή. 3. Ο κατηγορούμενος που είναι αλλοδαπός και στερείται την ελευθερία '
                              'του έχει δικαίωμα να ζητήσει να ενημερωθούν, χωρίς αδικαιολόγητη καθυστέρηση, οι '
                              'προξενικές αρχές του κράτους του οποίου είναι υπήκοος.' )

    main_paragraph7 = doc.add_paragraph ()
    main_paragraph7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph7.add_run (
        'Ε.  Εκ του άρθρου 98 του Κ.Π.Δ. , ήτοι δικαιωμα επικοινωνίας με τρίτα πρόσωπα κατά την '
        'διάρκεια της στέρησης της ελευθερίας . ' )

    main_paragraph8 = doc.add_paragraph ()
    main_paragraph8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph8.add_run ( '1.	Ο κατηγορούμενος που στερείται την ελευθερία του έχει δικαίωμα επικοινωνίας, χωρίς '
                              'αδικαιολόγητη καθυστέρηση, με ένα τουλάχιστον τρίτο πρόσωπο που έχει υποδείξει ο ίδιος. '
                              'Σε εξαιρετικές περιπτώσεις και για να αποτραπεί ο άμεσος κίνδυνος της παρ. 2 του'
                              ' προηγούμενου άρθρου, οι αρμόδιες αρχές μπορούν να περιορίζουν ή να αναβάλλουν την '
                              'άσκηση του ανωτέρω δικαιώματος. Στην περίπτωση αυτή εξετάζεται πρώτα αν ο κατηγορούμενος'
                              ' μπορεί να επικοινωνήσει με ένα άλλο πρόσωπο που αυτός υποδεικνύει. 2. Ο κατηγορούμενος'
                              ' που είναι αλλοδαπός και στερείται την ελευθερία του έχει δικαίωμα να επικοινωνεί, χωρίς '
                              'αδικαιολόγητη καθυστέρηση, με τις προξενικές αρχές του κράτους του οποίου είναι υπήκοος. '
                              'Έχει επίσης δικαίωμα επίσκεψης από τις προξενικές του αρχές, δικαίωμα συνομιλίας και '
                              'αλληλογραφίας μαζί τους και δικαίωμα να κανονίζεται η νομική του εκπροσώπηση από αυτές, '
                              'εφόσον οι εν λόγω αρχές δεν έχουν αντίρρηση.' )

    main_paragraph9 = doc.add_paragraph ()
    main_paragraph9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph9.add_run ( 'ΣΤ. Εκ του άρθρου 99 του Κ.Π.Δ., ήτοι δικαίωμα παράστασης του κατηγορουμένου με '
                              'συνήγορο :' )

    main_paragraph10 = doc.add_paragraph ()
    main_paragraph10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph10.add_run (
        '1. Ο κατηγορούμενος έχει το δικαίωμα στην απολογία του, ακόμη και σ’ αυτήν που γίνεται σε'
        ' αντιπαράσταση με μάρτυρες ή άλλους κατηγορουμένους, να παρίσταται με συνήγορο. Γι’ αυτό '
        'το σκοπό καλείται εικοσιτέσσερις ώρες πριν από κάθε ανακριτική ενέργεια. 2. Επιτρέπεται'
        ' σύντμηση της προθεσμίας αυτής, αν από την αναβολή δημιουργείται κίνδυνος που η ύπαρξή '
        'του βεβαιώνεται ειδικά με έκθεση του ανακριτή ή του ανακριτικού υπαλλήλου. 3. Ο ανακριτής'
        'έχει την υποχρέωση να διορίσει αυτεπαγγέλτως συνήγορο στον κατηγορούμενο για κακούργημα. '
        'Την ίδια υποχρέωση έχει και στα πλημμελήματα, αν το ζητήσει ρητά ο κατηγορούμενος. '
        '4. Σε καμιά περίπτωση δεν μπορεί να απαγορευθεί η επικοινωνία του κατηγορουμένου με τον '
        'συνήγορό του. Η επικοινωνία αυτή είναι απολύτως απόρρητη.' )

    main_paragraph11 = doc.add_paragraph ()
    main_paragraph11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph11.add_run ( 'Ζ.  Εκ του άρθρου 100 του Κ.Π.Δ., ήτοι δικαίωμα πρόσβασης στο υλικό της δικογραφίας:' )

    main_paragraph12 = doc.add_paragraph ()
    main_paragraph12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph12.add_run ( 'Ανακοίνωση των εγγράφων της ανάκρισης. 1. Ο ανακριτής, μόλις μετά την κλήτευσή του '
                               'εμφανισθεί ή οδηγηθεί σε αυτόν ο κατηγορούμενος για να απολογηθεί, του ανακοινώνει το '
                               'περιεχόμενο του κατηγορητηρίου και των εγγράφων της ανάκρισης. Επιτρέπεται επίσης στον'
                               ' κατηγορούμενο να μελετήσει ο ίδιος ή ο συνήγορός του το κατηγορητήριο και τα έγγραφα '
                               'της ανάκρισης. Με γραπτή αίτηση του κατηγορουμένου και με δαπάνη του χορηγούνται σε '
                               'αυτόν αντίγραφα του κατηγορητηρίου και των εγγράφων της ανάκρισης. 2. Την ίδια υποχρέωση'
                               ' έχει ο ανακριτής, και τα ίδια δικαιώματα ο κατηγορούμενος, όταν κληθεί ξανά σε'
                               ' συμπληρωματική απολογία. Πάντως μετά το τέλος της ανάκρισης και προτού διαβιβαστεί η '
                               'δικογραφία στον εισαγγελέα (άρθρο 308 παρ. 1), καλείται πάντοτε ο κατηγορούμενος να '
                               'μελετήσει όλη τη δικογραφία. Αν όμως η ανάκριση εξακολούθησε περισσότερο από μήνα μετά'
                               ' την πρώτη ή κάθε μεταγενέστερη απολογία, δικαιούται ο κατηγορούμενος να ασκεί τα '
                               'δικαιώματά του μια φορά το μήνα, και κάθε φορά ο ανακριτής συντάσσει σχετική έκθεση κάτω'
                               ' από την απολογία του κατηγορουμένου.' )

    main_paragraph13 = doc.add_paragraph ()
    main_paragraph13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph13.add_run ( 'Η.  Εκ του άρθρου 101 του Κ.Π.Δ. , ήτοι δικαίωμα διερμηνείας και μετάφρασης  :' )

    main_paragraph14 = doc.add_paragraph ()
    main_paragraph14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph14.add_run (
        'Ο ύποπτος ή ο κατηγορούμενος, που δεν ομιλεί ή δεν κατανοεί επαρκώς την ελληνική γλώσσα, '
        'έχει το δικαίωμα σε διερμηνεία καθώς και σε γραπτή μετάφραση όλων των ουσιωδών εγγράφων '
        'της διαδικασίας, σύμφωνα με όσα προβλέπονται στα άρθρα 233 και 237. ' )

    main_paragraph15 = doc.add_paragraph ()
    main_paragraph15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph15.add_run ( 'Θ. Εκ του άρθρου 103 του Κ.Π.Δ., ήτοι δικαίωμα προθεσμίας για την απολογία  :' )

    main_paragraph16 = doc.add_paragraph ()
    main_paragraph16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph16.add_run (
        '1. Ο κατηγορούμενος έχει το δικαίωμα να ζητήσει προθεσμία τουλάχιστον σαράντα οκτώ ωρών '
        'και δεν έχει υποχρέωση να απολογηθεί πριν περάσει η προθεσμία. 2. Ο ανακριτής μπορεί να '
        'παρατείνει την προθεσμία ύστερα από αίτηση του κατηγορουμένου.  ' )

    main_paragraph17 = doc.add_paragraph ()
    main_paragraph17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph17.add_run ( 'Ι.  Εκ του άρθρου 104 του Κ.Π.Δ. , ήτοι δικαίωμα σιωπής και μη αυτοενοχοποίησης :' )

    main_paragraph18 = doc.add_paragraph ()
    main_paragraph18.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph18.add_run ( '1. Ο ύποπτος ή ο κατηγορούμενος έχουν δικαίωμα σιωπής και μη αυτοενοχοποίησης. '
                               '2. Η άσκηση του δικαιώματος μη αυτοενοχοποίησης δεν εμποδίζει τη νόμιμη συγκέντρωση '
                               'αποδεικτικών στοιχείων, που υπάρχουν ανεξάρτητα από τη βούληση των υπόπτων και των '
                               'κατηγορουμένων.'
                               ' 3. Η άσκηση του δικαιώματος σιωπής και μη αυτοενοχοποίησης δεν μπορεί να αξιοποιηθεί σε '
                               'βάρος των υπόπτων και των κατηγορουμένων. ' )

    main_paragraph19 = doc.add_paragraph ()
    main_paragraph19.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph19.add_run ( 'ΙΑ.  Εκ του άρθρου 105 του Κ.Π.Δ. , ήτοι δικαιώματα στην αυτεπάγγελτη προανάκριση  :' )

    main_paragraph20 = doc.add_paragraph ()
    main_paragraph20.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph20.add_run (
        '1.	Όταν ενεργείται προανάκριση σύμφωνα με το άρθρο 245 παρ. 2, η εξέταση γίνεται όπως '
        'ορίζεται στις διατάξεις των άρθρων 273 και 274 και εκείνος που εξετάζεται έχει τα '
        'δικαιώματα που αναφέρονται στα άρθρα 91, 95, 96, 97, 98, 100, 101, 103 και 104. '
        '2. Ο εξεταζόμενος έχει το δικαίωμα πρόσβασης σε δικηγόρο χωρίς αδικαιολόγητη καθυστέρηση '
        'και σε κάθε περίπτωση προτού εξεταστεί από την αστυνομία ή άλλη αρχή επιβολής του νόμου '
        'ή δικαστική αρχή. Η διάταξη του άρθρου 99 παρ. 4 εφαρμόζεται και σε αυτή την περίπτωση. '
        '3. Η κατά παράβαση του παρόντος άρθρου εξέταση είναι άκυρη και δεν λαμβάνεται υπόψη.'
        ' Κατά τα άλλα εφαρμόζεται η παρ. 3 του άρθρου 244.' )

    main_paragraph21 = doc.add_paragraph ()
    main_paragraph21.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph21.add_run ( 'Κατόπιν της ανωτέρω γνωστοποίησης ο κατηγορούμενος δήλωσε ότι ενημερώθηκε'
                               ' πλήρως για τα παραπάνω δικαιώματα του και ερωτηθείς αν επιθυμεί να τα '
                               'ασκήσει απάντησε ότι δεν επιθυμεί να ασκήσει κάποιο από τα δικαιώματα '
                               'του ή ....... ' )

    main_paragraph22 = doc.add_paragraph ()
    main_paragraph22.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph22.add_run ( 'Η παρούσα έκθεση άρχισε να συντάσσεται την {{start_time_rights}}'
                               ' ώρα και περατώθηκε την {{end_time_rights}}  ώρα' )

    main_paragraph23 = doc.add_paragraph ()
    main_paragraph23.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph23.add_run ( 'Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού αναγνώσθηκε και βεβαιώθηκε,'
                               ' υπογράφεται ως ακολούθως:' )

    new_para5 = doc.add_paragraph ( "  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                    "         Ο   Ανακριτικός Υπάλληλος" )
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save ( 'sample_rights.docx' )
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ( "sample_rights.docx" )
    template.render ( context )
    template.save ( "ΔΙΚΑΙΩΜΑΤΑ_ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ.docx" )
    # Αποθήκευση του εγγράφου


def vas():
    global context  # Declare we're using the global context
    # Δημιουργία εγγράφου
    doc = Document ( )
    # Προσθήκη τίτλου (κεντραρισμένου)
    title = doc.add_heading ( 'ΕΓΓΡΑΦΗ ΣΤΟ ΒΙΒΛΙΟ ΑΔΙΚΗΜΑΤΩΝ ΚΑΙ ΣΥΜΒΑΝΤΩΝ' , level = 0 )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph = doc.add_paragraph ( )
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph.add_run ( 'Στις {{hour}} ώρα τις {{date_num}} {{month}} {{year }} προσήλθε στην Υπηρεσία μας ο'
                             '{{surname}} {{name}} του {{ fathername}} και της {{ mothername}} γεν. '
                             '{{dateOfBirth}} στη {{ placeOfBirth }} κατ. {{ address }}, αριθμός τηλεφώνου {{ tel }},'
                             'ηλεκτρονικό ταχυδρομείο {{ email }}, κάτοχος του υπ αριθμόν {{ DAT }} που εκδόθηκε την '
                             ' {{ issued }} από {{place_issued}} Α.Φ.Μ : {{ afm }}, Δ.Ο.Υ : {{ doy }}, και υπέβαλε'
                             'μήνυση εις βάρος του  {{surnamePerperator}} {{namePerperator}} του {{fathernamePerperator}}'
                             'και της {{ mothernamePerperator}} γεν. {{dateOfBirthPerperator }} στη '
                             '{{placeOfBirthPerperator }} κατ.{{ addressPerperator }},αριθμός τηλεφώνου '
                             '{{ telPreperator }}, ηλεκτρονικό ταχυδρομείου{{ emailPreperator }}, κάτοχος του υπ αριθμόν '
                             '{{ DATperperator }} που εκδόθηκε την {{ issuedPerperator }} από {{place_issuedPerperator}}'
                             'Α.Φ.Μ : {{ afmPreperator }}, Δ.Ο.Υ : {{ doyPrep }}, για παραβάσεις του/των άρθρων '
                             '{{offences  }} του Π.Κ πράξεις που έλαβαν χώρα  στη {{placeOfCrime}} στις {{ dateOfCrime}} '
                             'και περί ώρα {{ hourOfCrime }}. Δράστης συνελήφθη και με τη κατά αυτού σχηματισθείσα '
                             'δικογραφία οδηγηθεί κ. Εισαγγελέα Πλημ/κων Θεσσαλονίκης. Δόθηκε παραγγελία ιατροδικαστικής'
                             'εξέτασης.' )

    doc.save ( 'sample_vas.docx' )
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ( "sample_vas.docx" )
    template.render ( context )
    template.save ( "ΒΑΣ.docx" )
    # Αποθήκευση του εγγράφου

    # Κλήση της συνάρτησης


def transmission():
    global context  # Declare we're using the global context
    # Δημιουργία εγγράφου
    doc = Document ( )

    # Δημιουργία πίνακα με 1 γραμμή και 2 κελιά (για αριστερά & δεξιά κείμενο)
    table = doc.add_table ( rows = 1 ,cols = 2 )
    table.autofit = False  # Απενεργοποίηση αυτόματου προσαρμογής πλάτους

    # Ορισμός πλάτους στηλών (50% - 50%)
    col_widths = [Inches ( 3 ) ,Inches ( 3 )]
    for i ,width in enumerate ( col_widths ):
        table.columns[i].width = width

    # Κελί 1 (Αριστερά) - "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ"
    left_cell = table.cell ( 0 ,0 )
    left_texts = [
        "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ" ,
        "ΕΛΛΗΝΙΚΗ ΑΣΤΥΝΟΜΙΑ" ,
        "ΔΙΕΥΘ. ΑΣΤ. ΘΕΣΣΑΛΟΝΙΚΗΣ" ,
        "ΑΣΤΥΝΟΜΙΚΟ ΤΜΗΜΑ ΘΕΡΜΗΣ" ,
        "ΤΗΛ:2310461203" ,
        "email: atthermis@astynomia.gr" ,
        "Αρμόδιος:{{first_officer }}" ,
        "Αρ.πρωτ: {{ protocolnumber }}"
    ]

    for text in left_texts:
        p = left_cell.add_paragraph ( )
        run = p.add_run ( text )
        run.bold = True
        run.font.size = Pt ( 10 )  # γραμματοσειρά (12pt)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt ( 4 )

    # Κελί 2 (Δεξιά) - Τοποθεσία, ημερομηνία + επιπλέον κείμενο
    right_cell = table.cell ( 0 ,1 )

    # Γραμμή 1: Τοποθεσία & Ημερομηνία
    right_para1 = right_cell.add_paragraph ( )
    right_run1 = right_para1.add_run ( f"Θέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
    right_run1.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Γραμμή 2: "Π Ρ Ο Σ" (με κενά)
    right_para2 = right_cell.add_paragraph ( )
    right_run2 = right_para2.add_run ( " Π Ρ Ο Σ" )
    right_run2.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Γραμμή 3: Εργαστήριο & Πανεπιστήμιο
    right_para3 = right_cell.add_paragraph ( )
    right_run3 = right_para3.add_run (
        "ΕΙΣΑΓΓΕΛΙΑ ΠΛΗΜ/ΚΩΝ \nΘΕΣΣΑΛΟΝΙΚΗΣ\n" )
    right_run3.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph ( ).add_run ( ).add_break ( )  # Κενή γραμμή πριν το ΘΕΜΑ

    p_theme = doc.add_paragraph ( )
    run_theme = p_theme.add_run ( "ΘΕΜΑ: «Αποστολή Συνοδείας»" )
    run_theme.bold = True
    run_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    p_theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν

    main_theme = doc.add_paragraph ( )
    run1_theme = main_theme.add_run ( '   Αποστέλεται συνοδεία Αστυνομικών Υπηρεσίας μας ο {{surnamePerperator}} '
                                      '{{namePerperator}} του {{fathernamePerperator}} και της {{mothernamePerperator}}'
                                      'γεν. {{dateOfBirthPerperator }} στη {{placeOfBirthPerperator }} κατ.{'
                                      '{ addressPerperator }},αριθμός τηλεφώνου {{ telPreperator }}, ηλεκτρονικό '
                                      'ταχυδρομείου{{ emailPreperator }}, κάτοχος του υπ αριθμόν {{ DATperperator}} που'
                                      ' εκδόθηκε την {{ issuedPerperator }} από {{place_issuedPerperator}} Α.Φ.Μ : '
                                      '{{afmPreperator}}, Δ.Ο.Υ : {{ doyPrep }}, κατηγορούμενος για παραβάσεις του/των'
                                      'άρθρων {{ offences}} πράξεις που έλαβαν χώρα  στη {{placeOfCrime}} στις '
                                      '{{ dateOfCrime}} και περί ώρα {{ hourOfCrime }}' )
    run1_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση

    footer_para = doc.add_paragraph ( )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Προσθήκη κειμένου με ακριβή θέση (χρήση tabs)
    footer_run = footer_para.add_run ( )
    footer_run.add_text(f"\t\tΘέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
    footer_run.add_text("\t\t\tΟ\t\t\n")  # 5 tabs για στοίχιση
    footer_run.add_text("\t\t\t\t\tΑνακριτικός Υπάλληλος\n\n")
    footer_run.add_text("\t\t\t\t\t {{ first_officer }}\n")

    # Αποθήκευση
    doc.save('sample_transmission.docx')
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ('sample_transmission.docx')
    template.render(context)
    template.save("ΔΙΑΒΙΒΑΣΤΙΚΟ.docx")
    # Αποθήκευση του εγγράφου


def mutual():
    global context
    # Add mutual accusation specific data to the context
    doc = Document ()
    # Προσθήκη τίτλου (κεντραρισμένου)
    title = doc.add_heading ( 'ΕΚΘΕΣΗ ΕΝΟΡΚΗΣ ΕΞΕΤΑΣΗΣ' , level = 0 )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph = doc.add_paragraph ()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph.add_run (
        '    Στην {{ place1 }} σήμερα την {{ date_num1 }} του μήνα {{ month1 }} του έτους {{ year }}'
        ' ημέρα εβδομάδας '
        '{{ day1 }} και ώρα {{ hour1 }} ενώπιον εμού του {{ first_officer }} του {{ policeStation }} '
        'Θεσσαλονίκης, παρισταμένου και της {{ sec_officer }} της ιδίας υπηρεσίας, που προσλήφθηκε '
        'ως Β\' Ανακριτικός Υπάλληλος, εμφανίστηκε ο κατωτέρω μάρτυρας, ο οποίος αφού ρωτήθηκε για την ταυτότητα '
        'του κ.λ.π. απάντησε ότι ονομάζεται: {{surnamePerperator}} {{namePerperator}} του {{fathernamePerperator}}'
        ' και της {{mothernamePerperator}} γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ. '
        '{{addressPerperator}}, αριθμός τηλεφώνου {{ telPreperator}}, ηλεκτρονικό '
        'ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DATperperator}} που εκδόθηκε την  {{ issuedPerperator }}'
        ' από {{place_issuedPerperator}} Α.Φ.Μ : {{ afmPreperator}}, Δ.Ο.Υ : {{doyPrep}}'
    )
    # ΠΡΟΣΘΗΚΗ ΝΕΑΣ ΠΑΡΑΓΡΑΦΟΥ - Παράδειγμα 1: Απλή παράγραφος
    new_para1 = doc.add_paragraph ( "   Έπειτα ο εξεταζόμενος έδωσε τον προβλεπόμενο από τα άρθρα 219 και 220 παρ. 1 "
                                    " του Κώδικα Ποινικής Δικονομίας όρκο, ως ακολούθως: «Δηλώνω, επικαλούμενος την "
                                    "τιμή και την συνείδηση μου, ότι θα πω όλη την αλήθεια και μόνο την αλήθεια, "
                                    "χωρίς"
                                    " να προσθέσω ούτε να αποκρύψω τίποτα», και στην συνέχεια εξετάσθηκε ως εξής:." )
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph ( "ΕΡΩΤΗΣΗ: Ρωτήθηκε σχετικά:" )
    new_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para3 = doc.add_paragraph (
        "ΑΠΟΚΡΙΣΗ: Στις {{ dateOfCrime1 }} και περί ώρα {{ hourOfCrime1 }} στη {{ placeOfCrime1 }},"
        " {{ stateOfVictim1 }}, o {{surname}} {{name}} του {{ fathername}} και της {{ mothername}}"
        "γεν. {{dateOfBirth}} στη {{ placeOfBirth}} κατ.{{ address}},αριθμός τηλεφώνου {{ tel}},"
        " ηλεκτρονικό ταχυδρομείου{{ email}},κάτοχος του υπ αριθμόν {{ DAT }} που εκδόθηκε την  "
        "{{ issued}} από {{place_issued}} Α.Φ.Μ : {{ afm}}, Δ.Ο.Υ : {{ doy}}{{whatHappened1}},{{ howHappened1}}."
        "{{whyHappened1}}.{{add_something1}}. {{forensicExam1 }}.{{prosecution}}."
        "Τίποτε άλλο δεν έχω να προσθέσω και υπογράφω," )
    new_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para4 = doc.add_paragraph ( "Η παρούσα έκθεση άρχισε να συντάσσεται την  {{hour1}} ώρα και περαιώθηκε την "
                                    "{{hourOfReportFinished1}} ώρα. Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία "
                                    "αφού αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:" )
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5 = doc.add_paragraph ( "  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                    "         Ο   Ανακριτικός Υπάλληλος" )
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Αποθήκευση προτύπου
    doc.save ( 'sample_mutual.docx' )
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ( "sample_mutual.docx" )
    template.render ( context )
    template.save ( "Εκθεση_Ένορκης_Εξέτασης2.docx" )




def iatrodik_mutual():
    global context
    # Δημιουργία εγγράφου
    doc = Document ( )

    # Δημιουργία πίνακα με 1 γραμμή και 2 κελιά (για αριστερά & δεξιά κείμενο)
    table = doc.add_table(rows = 1, cols = 2)
    table.autofit = False  # Απενεργοποίηση αυτόματου προσαρμογής πλάτους

    # Ορισμός πλάτους στηλών (50% - 50%)
    col_widths = [Inches(3), Inches(3)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Κελί 1 (Αριστερά) - "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ"
    left_cell = table.cell(0, 0)
    left_texts = [
        "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ" ,
        "ΕΛΛΗΝΙΚΗ ΑΣΤΥΝΟΜΙΑ" ,
        "ΔΙΕΥΘ. ΑΣΤ. ΘΕΣΣΑΛΟΝΙΚΗΣ" ,
        "ΑΣΤΥΝΟΜΙΚΟ ΤΜΗΜΑ ΘΕΡΜΗΣ" ,
        "ΤΗΛ:2310461203" ,
        "email: atthermis@astynomia.gr" ,
        "Αρμόδιος:{{first_officer }}" ,
        "Αρ.πρωτ: {{ protocolnumber }}"
    ]

    for text in left_texts:
        p = left_cell.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)  # γραμματοσειρά (12pt)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)

    # Κελί 2 (Δεξιά) - Τοποθεσία, ημερομηνία + επιπλέον κείμενο
    right_cell = table.cell(0, 1)

    # Γραμμή 1: Τοποθεσία & Ημερομηνία
    right_para1 = right_cell.add_paragraph()
    right_run1 = right_para1.add_run(f"Θέρμη,{datetime.now().strftime('%d/%m/%Y')}")
    right_run1.font.size = Pt(12)  # γραμματοσειρά (12pt)
    right_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Γραμμή 2: "Π Ρ Ο Σ" (με κενά)
    right_para2 = right_cell.add_paragraph()
    right_run2 = right_para2.add_run(" Π Ρ Ο Σ")
    right_run2.font.size = Pt(12)  # γραμματοσειρά (12pt)
    right_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Γραμμή 3: Εργαστήριο & Πανεπιστήμιο
    right_para3 = right_cell.add_paragraph()
    right_run3 = right_para3.add_run (
        "ΕΡΓΑΣΤΗΡΙΟ ΙΑΤΡΟΔΙΚΑΣΤΙΚΗΣ ΚΑΙ\nΤΟΞΙΚΟΛΟΓΙΑΣ\nΑριστοτέλειο Πανεπιστήμιο Θεσσαλονίκης" )
    right_run3.font.size = Pt(12)  # γραμματοσειρά (12pt)
    right_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- ΚΕΝΤΡΙΚΟ ΚΕΙΜΕΝΟ (ΘΕΜΑ) ---
    doc.add_paragraph().add_run().add_break()  # Κενή γραμμή πριν το ΘΕΜΑ

    p_theme = doc.add_paragraph()
    run_theme = p_theme.add_run("ΘΕΜΑ: «Παραγγελία Ιατροδικαστικής εξέτασης»")
    run_theme.bold = True
    run_theme.font.size = Pt(14)  # γραμματοσειρά (14pt)
    p_theme.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_theme.paragraph_format.space_before = Pt(16)  # Επιπλέον κενό πριν

    main_theme = doc.add_paragraph()

    offences_text1 = ", ".join(context['offences1'][:-1])
    if len(context['offences1']) > 1:
        offences_text1 += f" and {context[ 'offences1' ][ -1 ]}"
    else :
        offences_text1 = context[ 'offences1' ][ 0 ]

    run1_theme = main_theme.add_run (
        "    Παρακαλούμε μεριμνήσετε για την εξέταση από Ιατροδικαστή της υπηρεσίας σας, "
        'του {{surnamePerperator}} {{namePerperator}} του {{fathernamePerperator}}'
        ' και της {{mothernamePerperator}} γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ. '
        '{{addressPerperator}}, αριθμός τηλεφώνου {{ telPreperator}}, ηλεκτρονικό '
        'ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DATperperator}} που εκδόθηκε την  {{ issuedPerperator }}'
        ' από {{place_issuedPerperator}} Α.Φ.Μ : {{ afmPreperator}}, Δ.Ο.Υ : {{doyPrep}}'
        ' διότι την {{ dateOfCrime1 }}  και ώρα {{ hourOfCrime1 }}επί της {{ placeOfCrime1}} '
        'δέχτηκε σωματικές βλάβες από τον {{surname}} {{name}} του {{ fathername}} και της {{ mothername}} γεν. '
        '{{dateOfBirth  }} στη {{ placeOfBirth }} κατ. {{ address }}, αριθμός τηλεφώνου '
        '{{ tel }}, ηλεκτρονικό ταχυδρομείο {{ email }},Α.Φ.Μ :{{ afm }}, Δ.Ο.Υ :{{ doy }} '
        "όπως ρητά δηλώθηκε στην από "
        " {{ date_num1 }} {{month1 }}  {{year1  }} και "
        "ώρα {{ hour1 }} έγκληση σε βάρος του προρριθέντα κατά παράβαση του  άρθρου  "
        "{{offences_text1 }}του Π.Κ. Συγκεκριμένα "
        "και όπως μας δήλωσε παθόντας προκλήθηκαν σωματικές βλάβες  {{ placeOfHurt1 }}." )
    run1_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    main_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν

    main2_theme = doc.add_paragraph ( )
    run2_theme = main2_theme.add_run (
        "Παρακαλούμε όπως με σχετική έκθεση σας μας απαντήσετε στα κατωτέρω ερωτήματα μας." )
    run2_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    main2_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main2_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν

    # --- ΑΡΙΘΜΗΜΕΝΗ ΛΙΣΤΑ ---
    list_items = [
        "Φέρει σωματικές  κακώσεις και ποιες" ,
        "Προ πόσου χρόνου προξενήθηκαν αυτές." ,
        "Με ποιο όργανο προξενήθηκαν." ,
        "Η σωματική κάκωση ή βλάβη είναι απλή." ,
        "Η σωματική βλάβη είναι επικίνδυνη δηλαδή τελέστηκε με τρόπο που θα μπορούσε να"
        " προκαλέσει κίνδυνο ζωής του "
        "παθόντος ή βαριά σωματική βλάβη." ,
        "Ότι άλλο  από  την επιστήμη σας δύνασθε να μας γνωρίσετε. "
    ]

    for item in list_items :
        p_list = doc.add_paragraph ( style = 'List Number' )  # Αριθμημένη λίστα
        p_list.add_run ( item ).font.size = Pt ( 14 )  # Μέγεθος γραμματοσειράς 14pt
        p_list.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Προαιρετικό: Justify

    main3_theme = doc.add_paragraph ( )
    run3_theme = main3_theme.add_run (
        "Την έκθεση πραγματογνωμοσύνης να μας την εγχειρίσετε το συντομότερο δυνατόν." )
    run3_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    main3_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main3_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν

    # --- ΚΑΤΩ ΔΕΞΙΑ ΠΛΗΡΟΦΟΡΙΑ ---
    # Προσθήκη κενού πριν την υπογραφή
    doc.add_paragraph ( ).add_run ( ).add_break ( )

    # Γραμμή 1: "ΕΚΤΕΛΕΣΘΗΚΕ ΚΑΙ ΚΑΤΑΤΕΘΗΚΕ..."
    p1_down = doc.add_paragraph ( )
    p1_down.add_run (
        "ΕΚΤΕΛΕΣΘΗΚΕ ΚΑΙ ΚΑΤΑΤΕΘΗΚΕ ΠΙΝΑΚΑΣ\nΙΑΤΡΟΔΙΚΑΣΤΙΚΩΝ ΔΙΚΑΙΩΜΑΤΩΝ\n\tΑπό Ευρώ.............................."
        ".........." ).font.size = Pt ( 14 )
    p1_down.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Γραμμή 2: Παύλα "-Ο-" (στο κέντρο)
    p2_down = doc.add_paragraph ( )
    p2_down.add_run (
        "                              -Ο-                                         \t\t-O-" ).font.size = Pt (
        14 )
    p2_down.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Γραμμή 3: "Ανακριτικός Υπάλληλος" (δεξιά)
    p3_down = doc.add_paragraph ( )
    p3_down.add_run ( "\tΑνακριτικός Υπάλληλος\t\t\tΑνακριτικός Υπάλληλος" ).font.size = Pt ( 14 )
    p3_down.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Ορισμός γραμματοσειράς (Arial, 12pt) για όλο το κείμενο
    for row in table.rows :
        for cell in row.cells :
            for paragraph in cell.paragraphs :
                for run in paragraph.runs :
                    run.font.name = 'Arial'

        # Αποθήκευση προτύπου
        doc.save ( 'sampleiatr_document.docx' )
        # Επεξεργασία και τελική αποθήκευση
        template = DocxTemplate ( "sampleiatr_document.docx" )
        template.render ( context )
        template.save ( "iatrodikastikh_mutual.docx" )
        # Αποθήκευση του εγγράφου


def silipsi_mutual():
    global context
    # Δημιουργία εγγράφου
    doc = Document()

    title = doc.add_heading ( 'ΕΚΘΕΣΗ ΣΥΛΛΗΨΗΣ' , level = 0 )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph = doc.add_paragraph ( )
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    offences_text = ", ".join ( context[ 'offences1' ][ :-1 ] )
    if len ( context[ 'offences1' ] ) > 1:
        offences_text += f" and {context[ 'offences1' ][ -1 ]}"
    else:
        offences_text = context[ 'offences1' ][ 0 ]

    main_paragraph.style.font.size = Pt ( 14 )
    main_paragraph.paragraph_format.space_before = Pt ( 16 )
    main_paragraph.add_run (
        '    Στην {{ place1 }} σήμερα την {{ date_num1 }} του μήνα {{ month1 }} του έτους {{ year1 }}'
        ' ημέρα εβδομάδας '
        '{{ day1}} και ώρα {{startTimeArrestReport1 }} ενώπιον εμού του {{ first_officer }} του {{ policeStation}}'
        'Θεσσαλονίκης, παρισταμένου και της {{ sec_officer }} της ιδίας υπηρεσίας, που προσλήφθηκε '
        'ως Β\' Ανακριτικός Υπάλληλος, οδηγήθηκε στο κατάστημα του Α.Τ. Θέρμης, ο κάτωθι {{surnamePerperator}} '
        ' {{namePerperator}} του {{ fathernamePerperator}} και της {{ mothernamePerperator}}'
        ' γεν. {{dateOfBirthPerperator }} στη {{ placeOfBirthPerperator }} κατ.'
        " {{ addressPerperator }},αριθμός τηλεφώνου {{ telPreperator }}, ηλεκτρονικό "
        " ταχυδρομείου{{ emailPreperator }}, κάτοχος του υπ αριθμόν {{ DATperperator }} που εκδόθηκε την  "
        "{{ issuedPerperator }} από {{place_issuedPerperator}}Α.Φ.Μ : {{ afmPreperator }}, Δ.Ο.Υ : {{ doyPrep }}, "
        "από τον {{officer_arrest1 }} υπηρετών στο  {{  policeStation1}}, που τον συνέλαβε στις {{ dateOfArrest1}} "
        "κι ώρα {{ hourOfArrest1}} στη {{ placeOfArrest1 }} για παράβαση του άρθρου {{offences_text1}} του Π.Κ."
    )

    main2_theme = doc.add_paragraph ( )
    run2_theme = main2_theme.add_run (
        "  Αφού εξέτασα αυτήν, διέταξα την παραπομπή και παράδοση της με την παρούσα έκθεση στον  κ. Εισαγγελέα"
        " Πλημμελειοδικών Θεσσαλονίκης." )
    run2_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    main2_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main2_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν

    main3_theme = doc.add_paragraph ( )
    run3_theme = main3_theme.add_run ( 'Η παρούσα έκθεση άρχισε να συντάσσεται την {{ startTimeArrestReport1}} ’ '
                                       'ώρα και περατώθηκε την {{ endTimeOfReport1}} ώρα.  '
                                       )
    run3_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    main3_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main3_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν

    new_para4 = doc.add_paragraph ( " Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού"
                                    "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:" )
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5 = doc.add_paragraph ( "  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                    "         Ο   Ανακριτικός Υπάλληλος" )
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Αποθήκευση προτύπου
    doc.save ( 'sampleArrest_document.docx' )
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ( "sampleArrest_document.docx" )
    template.render ( context )
    template.save ("ΕΚΘΕΣΗ_ΣΥΛΛΗΨΗΣ_mutual.docx")
    # Αποθήκευση του εγγράφου


def deltio_kat_mutual():
    global context
    # Δημιουργία εγγράφου
    doc = Document ( )
    title = doc.add_heading ( 'ΔΕΛΤΙΟ ΣΤΟΙΧΕΙΩΝ ΤΑΥΤΟΤΗΤΑΣ ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ' , level = 0 )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ορισμός γραμματοσειράς (bold + υπογραμμισμένο) για τον τίτλο
    for run in title.runs:
        run.font.size = Pt(14)
        run.bold = True  # Bold
        run.underline = WD_UNDERLINE.SINGLE  # Υπογράμμιση

        # ΕΠΩΝΥΜΟ
        p_eponymo = doc.add_paragraph( )
        run_label = p_eponymo.add_run("ΕΠΩΝΥΜΟ: ")
        run_label.bold = True
        run_value = p_eponymo.add_run ( " {{ surname }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝΟΜΑ
        p_onoma = doc.add_paragraph()
        run_label = p_onoma.add_run("ΟΝΟΜΑ:    ")
        run_label.bold = True
        run_value = p_onoma.add_run("{{ name }}")  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝΟΜΑ ΠΑΤΕΡΑ
        p_onoma_patera = doc.add_paragraph ( )
        run_label = p_onoma_patera.add_run ( "ΟΝΟΜΑ ΠΑΤΕΡΑ:    " )
        run_label.bold = True
        run_value = p_onoma_patera.add_run ( "{{ fathername}}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝΟΜΑ ΜΗΤΕΡΑΣ
        p_onoma_mhteras = doc.add_paragraph ( )
        run_label = p_onoma_mhteras.add_run ( "ΟΝΟΜΑ ΜΗΤΕΡΑΣ:    " )
        run_label.bold = True
        run_value = p_onoma_mhteras.add_run ( " {{ mothername}}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΠΑΤΡΙΚΟ ΕΠΩΝΥΜΟ (Για έγγαμες)
        p_onoma_patriko = doc.add_paragraph ( )
        run_label = p_onoma_patriko.add_run ( "ΠΑΤΡΙΚΟ ΕΠΩΝΥΜΟ (Για έγγαμεσ):    " )
        run_label.bold = True
        run_value = p_onoma_patriko.add_run ( "_______________________" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΟΝ/ΜΟ ΣΥΖΥΓΟΥ
        p_onoma_syzygou = doc.add_paragraph ( )
        run_label = p_onoma_syzygou.add_run ( "ΟΝ/ΜΟ ΣΥΖΥΓΟΥ:    " )
        run_label.bold = True
        run_value = p_onoma_syzygou.add_run ( "_______________________" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΗΜΕΡΟΜΗΝΙΑ ΓΕΝΝΗΣΗΣ
        p_dateofbirth = doc.add_paragraph ( )
        run_label = p_dateofbirth.add_run ( "ΗΜΕΡΟΜΗΝΙΑ ΓΕΝΝΗΣΗΣ:    " )
        run_label.bold = True
        run_value = p_dateofbirth.add_run ( "  {{ dateOfBirth  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΤΟΠΟΘΕΣΙΑ ΓΕΝΝΗΣΗΣ
        p_placeofbirth = doc.add_paragraph ( )
        run_label = p_placeofbirth.add_run ( " ΤΟΠΟΘΕΣΙΑ ΓΕΝΝΗΣΗΣ:    " )
        run_label.bold = True
        run_value = p_placeofbirth.add_run ( "  {{ placeOfBirth  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΤΟΠΟΘΕΣΙΑ κατοικιασ
        p_address = doc.add_paragraph ( )
        run_label = p_address.add_run ( " ΤΟΠΟΣ ΚΑΤΟΙΚΙΑΣ:    " )
        run_label.bold = True
        run_value = p_address.add_run ( "  {{ address }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # Δ.Α.Τ
        p_dat = doc.add_paragraph ( )
        run_label = p_dat.add_run ( " ΑΡΙΘΜΟΣ Δ.Α.Τ:    " )
        run_label.bold = True
        run_value = p_dat.add_run ( "  {{ DAT }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # Δ.Α.Τ
        p_issue = doc.add_paragraph ( )
        run_label = p_issue.add_run ( " ΕΚΔΟΘΕΝ:    " )
        run_label.bold = True
        run_value = p_issue.add_run ( "  {{ issued  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # Δ.Α.Τ
        p_issueplace = doc.add_paragraph ( )
        run_label = p_issueplace.add_run ( " ΑΠΟ:    " )
        run_label.bold = True
        run_value = p_issueplace.add_run ( " {{ place_issued  }}" )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        # ΑΦΜ
        # Δ.Α.Τ
        p_afm = doc.add_paragraph ( )
        run_label = p_afm.add_run ( " Α.Φ.Μ:    " )
        run_label.bold = True
        run_value = p_afm.add_run ( "  {{ afm}} από Δ.Ο.Υ {{  doy }}  " )  # Υπογραμμισμένο κενό
        run_value.underline = WD_UNDERLINE.SINGLE

        footer_para = doc.add_paragraph ( )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Προσθήκη κειμένου με ακριβή θέση (χρήση tabs)
        footer_run = footer_para.add_run ( )
        footer_run.add_text ( f"\t\tΘέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
        footer_run.add_text ( "\t\t\tΟ\t\t\n" )  # 5 tabs για στοίχιση
        footer_run.add_text ( "\t\t\t\t\tΑνακριτικός Υπάλληλος\n\n" )
        footer_run.add_text ( "\t\t\t\t\t {{ first_officer }}\n" )

        # Αποθήκευση
        doc.save ( 'deltio_katigoroumenoutemplate.docx' )
        # Επεξεργασία και τελική αποθήκευση
        template = DocxTemplate ( 'deltio_katigoroumenoutemplate.docx' )
        template.render ( context )
        template.save ( "ΔΕΛΤΙΟ_ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ.docx" )
        # Αποθήκευση του εγγράφου



def apologia_mutual():
    global context
    # Δημιουργία εγγράφου
    doc = Document ( )
    # Προσθήκη τίτλου (κεντραρισμένου)
    title = doc.add_heading ( 'ΕΚΘΕΣΗ ΕΞΕΤΑΣΗΣ ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ' ,level = 0 )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph = doc.add_paragraph ( )
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph.add_run ( 'Στην {{ place1 }} σήμερα την {{ date_num_apologia1}} του μήνα {{ month_apologia1 }} του'
                             ' έτους {{year_apologia1}} ημέρα εβδομάδας {{day_apologia1}} και ώρα'
                             ' {{start_hour_apologia1}}'
                             ' ενώπιον εμού του {{ first_officer }} του {{ policeStation }} Θεσσαλονίκης, '
                             'παρισταμένου  '
                             'και του  {{ sec_officer }} της ιδίας υπηρεσίας, που προσλήφθηκε ως Β\' Ανακριτικός'
                             ' Υπάλληλος, εξετάζεται ο κατωτέρω σημειούμενος κατηγορούμενος ,' )

    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph1 = doc.add_paragraph ( )
    main_paragraph1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph1.add_run ( 'ΕΡΩΤΗΣΗ : Πως ονομάζεσαι κ.λ.π.' )

    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph2 = doc.add_paragraph ( )
    main_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph2.add_run (
        'ΑΠΟΚΡΙΣΗ : {{surname}} {{name}} του {{ fathername}} και της {{ mothername}} γεν.'
        ' {{dateOfBirth  }} στη {{ placeOfBirth }} κατ. {{ address }}, αριθμός τηλεφώνου '
        '{{ tel }}, ηλεκτρονικό ταχυδρομείο {{ email }}, κάτοχος του υπ αριθμόν {{ DAT }} που '
        'εκδόθηκε την  {{ issued }} από {{place_issued}} Α.Φ.Μ : {{ afm }}, Δ.Ο.Υ : {{ doy }}'
        ', από τον {{officer_arrest }} υπηρετών στο  {{  policeStation }}' )

    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph3 = doc.add_paragraph ( )
    main_paragraph3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph3.add_run ( 'Ενταύθα γνωρίσαμε στον εξεταζόμενο ότι κατηγορείται για παράβαση του άρθρου '
                              '{{offences_text1}} του Π.Κ.' )

    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph4 = doc.add_paragraph ( )
    main_paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph4.add_run ( 'Ειδικότερα κατηγορείται ότι στις {{dateOfCrime1}}και περί ώρα {{ hourOfCrime1}}'
                              ' στο {{placeOfCrime1}}  προέβης στην παράβαση του/των άρθρων {{offences1 }} τπυ Π.Κ.'
                              'εις βάρος του {{surnamePerperator}}  {{namePerperator}} του {{fathernamePerperator}} '
                              'και της {{ mothernamePerperator}} γεν. {{dateOfBirthPerperator }} στη '
                              '{{ placeOfBirthPerperator }}κατ.{{ addressPerperator }},αριθμός τηλεφώνου '
                              '{{ telPreperator }}, ηλεκτρονικό '
                              ' ταχυδρομείου{{ emailPreperator }}, κάτοχος του υπ αριθμόν {{ DATperperator }} που '
                              'εκδόθηκε την {{ issuedPerperator }} από {{place_issuedPerperator}}Α.Φ.Μ : '
                              '{{ afmPreperator }}, Δ.Ο.Υ : {{ doyPrep }}'
                              )

    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph5 = doc.add_paragraph ( )
    main_paragraph5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph5.add_run (
        'Στη συνέχεια αφού ανακοινώσαμε στον εξεταζόμενο το περιεχόμενο των εγγράφων της ανάκρισης,'
        ' εξηγήσαμε εκ νέου με σαφήνεια σ’ αυτόν , σύμφωνα με τα άρθρα 95 και 105 ή 106 του Κώδικα '
        'Ποινικής Δικονομίας , όλα τα δικαιώματά του που προβλέπονται από τα άρθρα  91,95,96,97,98,'
        '100,101 , 104 και 273 του Κώδικα Ποινικής Δικονομίας και ειδικότερα ότι δικαιούται : να '
        'παρίσταται μετά συνηγόρου , μετά του οποίου δεν μπορεί να απαγορευτεί η απολύτως απόρρητη '
        'επικοινωνία του σε καμιά περίπτωση, να μελετήσει ο ίδιος ή ο συνήγορός του  τα έγγραφα της'
        ' ανάκρισης και του κατηγορητηρίου (πραγματικών περιστατικών- κατηγορίας) , να του '
        'χορηγηθούν αντίγραφα αυτών με δική του δαπάνη και μετά από γραπτή αίτησή του να ζητήσει '
        'προθεσμία μέχρι 48 ωρών , προ της παρέλευσης της οποίας δεν υποχρεούται σε απολογία και '
        'ότι δύναται να δοθεί παράταση της προθεσμίας αυτής με αίτησή του , να αρνηθεί να απαντήσει'
        ' (δικαίωμα σιωπής και μη αυτοενεχοποίσης) , και να παραδώσει την απολογία του γραπτή . ' )

    main_paragraph6 = doc.add_paragraph ( )
    main_paragraph6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph6.add_run ( 'ΕΡΩΤΗΣΗ: Επιθυμείται να κάνετε χρήση των δικαιωμάτων που σας γνωστοποιήθηκαν ;  ' )

    main_paragraph6 = doc.add_paragraph ( )
    main_paragraph6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph6.add_run ( 'ΑΠΟΚΡΙΣΗ: {{ question_rights_apologia }}  ' )

    main_paragraph7 = doc.add_paragraph ( )
    main_paragraph7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph7.add_run ( 'ΑΠΟΚΡΙΣΗ: {{ question_rights_apologia }}  ' )

    main_paragraph8 = doc.add_paragraph ( )
    main_paragraph8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph8.add_run ( 'Ύστερα από τα ανωτέρω προβήκαμε στην λήψη της απολογίας του  ' )

    main_paragraph9 = doc.add_paragraph ( )
    main_paragraph9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph9.add_run ( 'Για πιστοποίηση συντάχθηκε  η παρούσα έκθεση, η οποία αφού αναγνώσθηκε και βεβαιώθηκε '
                              ' υπογράφεται ως ακολούθως : ' )

    main_paragraph10 = doc.add_paragraph ( )
    main_paragraph10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph10.add_run ( 'Ο Εξετασθείς                              Ο Β Ανακριτικός Υπάλληλος '
                               '              Ο   Ανακριτικός Υπάλληλος\n\n\n' )

    main_paragraph11 = doc.add_paragraph ( )
    main_paragraph11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph11.add_run ( ' Στη συνέχεια, προβήκαμε στην εξέταση του κατηγορουμένου ως ακολούθως : ' )

    main_paragraph12 = doc.add_paragraph ( )
    main_paragraph12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph12.add_run ( ' ΕΡΩΤΗΣΗ :  Κατηγορήθηκες  άλλη φορά και για ποια αιτία ;' )

    main_paragraph13 = doc.add_paragraph ( )
    main_paragraph13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph13.add_run ( 'ΑΠΟΚΡΙΣΗ: {{ question_crime_apologia }}  ' )

    main_paragraph14 = doc.add_paragraph ( )
    main_paragraph14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph14.add_run ( 'ΕΡΩΤΗΣΗ : Κατηγορείσαι ήδη για τις πράξεις που σου γνωστοποιήθηκαν ανωτέρω. '
                               'Τι απολογείσαι;   ' )

    main_paragraph15 = doc.add_paragraph ( )
    main_paragraph15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph15.add_run ( 'ΑΠΟΚΡΙΣΗ : {{question_crime_confess }} ' )

    new_para4 = doc.add_paragraph ( "Η παρούσα έκθεση άρχισε να συντάσσεται την  {{start_hour_apologia1}} ώρα και "
                                    "περαιώθηκε την {{end_hour_apologia1}} ώρα. Για πίστωση συντάχθηκε η παρούσα "
                                    "έκθεση η οποία αφού αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:" )
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    new_para5 = doc.add_paragraph ( "  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                    "         Ο   Ανακριτικός Υπάλληλος" )
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save ( 'sample_confess.docx' )
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ( "sample_confess.docx" )
    template.render ( context )
    template.save ( "ΕΚΘΕΣΗ_ΕΞΕΤΑΣΗΣ_ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ_mutual.docx" )
    # Αποθήκευση του εγγράφου


def rights_mutual(context):
    # Δημιουργία εγγράφου
    doc = Document ( )
    # Προσθήκη τίτλου (κεντραρισμένου)
    title = doc.add_heading ( 'ΕΚΘΕΣΗ ΕΞΕΤΑΣΗΣ ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ',level = 0 )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Προσθήκη παραγράφου με πλήρη στοίχιση (justify)
    main_paragraph = doc.add_paragraph ( )
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph.add_run ( 'Στην {{ plac1e }} σήμερα την {{ date_num_apologia}} του μήνα {{ month_apologia }} του'
                             ' έτους {{year_apologia}} ημέρα εβδομάδας {{day_apologia}} και ώρα {{start_time_rights}}'
                             ' ενώπιον εμού του {{first_officer }} του {{ policeStation }} Θεσσαλονίκης, παρισταμένου'
                             'και του  {{ sec_officer }} της ιδίας υπηρεσίας, που προσλήφθηκε ως Β\' Ανακριτικός'
                             ' Υπάλληλος, προσκλήθηκε ο  κατωτέρω σημειούμενος  κατηγορούμενος, που ονομάζεται '
                             '{{surname}} {{name}} του {{ fathername}} και της {{ mothername}} γεν.'
                             ' {{dateOfBirth  }} στη {{ placeOfBirth }} κατ. {{ address }}, αριθμός τηλεφώνου '
                             '{{ tel }}, ηλεκτρονικό ταχυδρομείο {{ email }}, κάτοχος του υπ αριθμόν {{ DAT }} που '
                             'εκδόθηκε την  {{ issued }} από {{place_issued}} Α.Φ.Μ : {{ afm }}, Δ.Ο.Υ : {{ doy }}'
                             ',στον οποίο γνωστοποιήσαμε ότι κατηγορείται '
                             'για παράβαση του/των άρθρων  {{offences1 }} τπυ Π.Κ. και εξηγήσαμε με σαφήνεια και'
                             ' πληρότητα σ’ αυτόν βάσει του άρθρου 95 του Κώδικα Ποινικής  Δικονομίας όλα τα εκ των '
                             ' άρθρων 91,95,96,97,98,99,100,103 και 104  του  Κ.Π.Δ.  δικαιώματά  του/της  και  '
                             'αναλυτικότερα :' )

    main_paragraph1 = doc.add_paragraph ( )
    main_paragraph1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph1.add_run ( 'Α. Εκ του άρθρου 91 του Κ.Π.Δ. , ήτοι δικαίωμα παροχής δωρεάν νομικής βοήθεια :' )

    main_paragraph2 = doc.add_paragraph ( )
    main_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph2.add_run (
        'Ο ύποπτος ή ο κατηγορούμενος έχουν δικαίωμα δωρεάν νομικής βοήθειας, το οποίο περιλαμβάνει'
        ' παροχή νομικών συμβουλών και νομική αρωγή και εκπροσώπηση τους ενώπιον του δικαστηρίου , '
        'σύμφωνα με όσα ορίζονται σε σχετικές διατάξεις. ' )

    main_paragraph2 = doc.add_paragraph ( )
    main_paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph2.add_run ( 'B. Εκ του άρθρου 95 του Κ.Π.Δ. , ήτοι δικαίωμα σε ενημέρωση :' )

    main_paragraph3 = doc.add_paragraph ( )
    main_paragraph3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph3.add_run (
        '1.	Ο ύποπτος ή ο κατηγορούμενος ενημερώνεται αμέσως όσον αφορά τουλάχιστον στα ακόλουθα '
        'δικαιώματα: α) το δικαίωμα παράστασης με συνήγορο, β) το δικαίωμα και τις προϋποθέσεις '
        'παροχής δωρεάν νομικών συμβουλών, γ) το δικαίωμα ενημέρωσης σχετικά με την κατηγορία, '
        'δ) το δικαίωμα διερμηνείας και μετάφρασης και ε) το δικαίωμα σιωπής και μη '
        'αυτοενοχοποίησης. 2. Η ενημέρωση σύμφωνα με την παρ. 1 παρέχεται σε απλή και κατανοητή '
        'γλώσσα, προφορικώς ή εγγράφως, λαμβανομένων υπόψη των ειδικών αναγκών των υπόπτων ή '
        'κατηγορουμένων που είναι ευάλωτα πρόσωπα. Αντικείμενο της ενημέρωσης οφείλει να αποτελεί '
        'και η αναφορά των συνεπειών παραίτησης από την άσκηση των δικαιωμάτων. Για την ενημέρωση '
        'και την απάντηση του υπόπτου ή του κατηγορουμένου συντάσσεται έκθεση, η οποία και '
        'υπογράφεται.' )

    main_paragraph4 = doc.add_paragraph ( )
    main_paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph4.add_run ( 'Γ. Εκ του άρθρου 96 , ήτοι δικαίωμα χορήγησης εγγράφου περί των δικαιωμάτων : ' )

    main_paragraph5 = doc.add_paragraph ( )
    main_paragraph5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph5.add_run (
        '1.	Στον ύποπτο ή τον κατηγορούμενο, ο οποίος συλλαμβάνεται ή κρατείται, παρέχεται αμέσως '
        'έγγραφο στο οποίο καταγράφονται τα δικαιώματά του και του επιτρέπεται να το διατηρεί στην'
        ' κατοχή του καθ’ όλη τη διάρκεια της στέρησης της ελευθερίας του. Το έγγραφο αυτό '
        'περιλαμβάνει πληροφορίες σχετικά με τα ακόλουθα δικαιώματα: α) το δικαίωμα παράστασης με '
        'συνήγορο, β) το δικαίωμα και τις προϋποθέσεις παροχής δωρεάν νομικών συμβουλών, γ) το '
        'δικαίωμα ενημέρωσης σχετικά με την κατηγορία, δ) το δικαίωμα διερμηνείας και μετάφρασης, '
        'ε) το δικαίωμα σιωπής και μη αυτοενοχοποίησης, στ) το δικαίωμα πρόσβασης στο υλικό της'
        ' δικογραφίας, ζ) το δικαίωμα ενημέρωσης των προξενικών αρχών και ενός επιπλέον προσώπου '
        'της επιλογής του, η) το δικαίωμα σε επείγουσα ιατρική περίθαλψη, θ) τον ανώτατο αριθμό'
        ' ωρών ή ημερών κατά τις οποίες ο κατηγορούμενος δύναται να στερηθεί της ελευθερίας του '
        'προτού προσαχθεί ενώπιον δικαστικής αρχής και ι) πληροφορίες σχετικά με τις δυνατότητες '
        'προσβολής του νόμιμου χαρακτήρα της σύλληψης ή της κράτησης. 2. Το ως άνω έγγραφο '
        'συντάσσεται σε απλή και κατανοητή γλώσσα. Όταν αυτό δεν είναι διαθέσιμο στην κατάλληλη '
        'γλώσσα, ο ύποπτος ή ο κατηγορούμενος ενημερώνεται για τα δικαιώματά του προφορικά σε '
        'γλώσσα που κατανοεί. Το εν λόγω έγγραφο πρέπει στη συνέχεια να χορηγείται, χωρίς'
        ' αδικαιολόγητη καθυστέρηση, σε γλώσσα που ο ύποπτος ή ο κατηγορούμενος κατανοεί.' )

    main_paragraph5 = doc.add_paragraph ( )
    main_paragraph5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph5.add_run (
        'Δ. Εκ του άρθρου 97 , ήτοι δικαιώμα ενημέρωσης προσώπου της επιλογής του κατηγορουμένου '
        'σε περίπτωση στέρησης της ελευθερίας του : ' )

    main_paragraph6 = doc.add_paragraph ( )
    main_paragraph6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph6.add_run ( '1. Ο κατηγορούμενος έχει δικαίωμα να ζητήσει να ενημερωθεί, χωρίς αδικαιολόγητη'
                              ' καθυστέρηση, για τη στέρηση της ελευθερίας του ένα τουλάχιστον πρόσωπο της επιλογής του. '
                              'Αν ο κατηγορούμενος είναι ανήλικος ενημερώνεται ο ασκών τη γονική μέριμνα, εκτός αν αυτό '
                              'αντιβαίνει στο συμφέρον του ανηλίκου, οπότε ενημερώνεται κάποιο άλλο ενδεδειγμένο ενήλικο '
                              'πρόσωπο ή η αρμόδια για την προστασία των ανηλίκων αρχή. 2. Σε εξαιρετικές περιπτώσεις και'
                              'για να αποτραπεί άμεσος κίνδυνος είτε για τη ζωή, την ελευθερία ή τη σωματική ακεραιότητα '
                              'κάποιου προσώπου είτε για την αποδεικτική διερεύνηση του εγκλήματος, οι αρμόδιες αρχές '
                              'μπορούν να μην επιτρέπουν προσωρινά την ενημέρωση τρίτου προσώπου για τη στέρηση της '
                              'ελευθερίας του κατηγορουμένου. Στην περίπτωση αυτή εξετάζεται αν ένα άλλο τρίτο πρόσωπο, '
                              'που έχει υποδειχθεί από τον κατηγορούμενο, μπορεί να ενημερωθεί σχετικά. Αν ο'
                              ' κατηγορούμενος είναι ανήλικος, η αρμόδια για την προστασία των ανηλίκων αρχή ενημερώνεται'
                              'στην περίπτωση αυτή. 3. Ο κατηγορούμενος που είναι αλλοδαπός και στερείται την ελευθερία '
                              'του έχει δικαίωμα να ζητήσει να ενημερωθούν, χωρίς αδικαιολόγητη καθυστέρηση, οι '
                              'προξενικές αρχές του κράτους του οποίου είναι υπήκοος.' )

    main_paragraph7 = doc.add_paragraph ( )
    main_paragraph7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph7.add_run (
        'Ε.  Εκ του άρθρου 98 του Κ.Π.Δ. , ήτοι δικαιωμα επικοινωνίας με τρίτα πρόσωπα κατά την '
        'διάρκεια της στέρησης της ελευθερίας . ' )

    main_paragraph8 = doc.add_paragraph ( )
    main_paragraph8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph8.add_run ( '1.	Ο κατηγορούμενος που στερείται την ελευθερία του έχει δικαίωμα επικοινωνίας, χωρίς '
                              'αδικαιολόγητη καθυστέρηση, με ένα τουλάχιστον τρίτο πρόσωπο που έχει υποδείξει ο ίδιος. '
                              'Σε εξαιρετικές περιπτώσεις και για να αποτραπεί ο άμεσος κίνδυνος της παρ. 2 του'
                              ' προηγούμενου άρθρου, οι αρμόδιες αρχές μπορούν να περιορίζουν ή να αναβάλλουν την '
                              'άσκηση του ανωτέρω δικαιώματος. Στην περίπτωση αυτή εξετάζεται πρώτα αν ο κατηγορούμενος'
                              ' μπορεί να επικοινωνήσει με ένα άλλο πρόσωπο που αυτός υποδεικνύει. 2. Ο κατηγορούμενος'
                              ' που είναι αλλοδαπός και στερείται την ελευθερία του έχει δικαίωμα να επικοινωνεί, χωρίς '
                              'αδικαιολόγητη καθυστέρηση, με τις προξενικές αρχές του κράτους του οποίου είναι υπήκοος. '
                              'Έχει επίσης δικαίωμα επίσκεψης από τις προξενικές του αρχές, δικαίωμα συνομιλίας και '
                              'αλληλογραφίας μαζί τους και δικαίωμα να κανονίζεται η νομική του εκπροσώπηση από αυτές, '
                              'εφόσον οι εν λόγω αρχές δεν έχουν αντίρρηση.' )

    main_paragraph9 = doc.add_paragraph ( )
    main_paragraph9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph9.add_run ( 'ΣΤ. Εκ του άρθρου 99 του Κ.Π.Δ., ήτοι δικαίωμα παράστασης του κατηγορουμένου με '
                              'συνήγορο :' )

    main_paragraph10 = doc.add_paragraph ( )
    main_paragraph10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph10.add_run (
        '1. Ο κατηγορούμενος έχει το δικαίωμα στην απολογία του, ακόμη και σ’ αυτήν που γίνεται σε'
        ' αντιπαράσταση με μάρτυρες ή άλλους κατηγορουμένους, να παρίσταται με συνήγορο. Γι’ αυτό '
        'το σκοπό καλείται εικοσιτέσσερις ώρες πριν από κάθε ανακριτική ενέργεια. 2. Επιτρέπεται'
        ' σύντμηση της προθεσμίας αυτής, αν από την αναβολή δημιουργείται κίνδυνος που η ύπαρξή '
        'του βεβαιώνεται ειδικά με έκθεση του ανακριτή ή του ανακριτικού υπαλλήλου. 3. Ο ανακριτής'
        'έχει την υποχρέωση να διορίσει αυτεπαγγέλτως συνήγορο στον κατηγορούμενο για κακούργημα. '
        'Την ίδια υποχρέωση έχει και στα πλημμελήματα, αν το ζητήσει ρητά ο κατηγορούμενος. '
        '4. Σε καμιά περίπτωση δεν μπορεί να απαγορευθεί η επικοινωνία του κατηγορουμένου με τον '
        'συνήγορό του. Η επικοινωνία αυτή είναι απολύτως απόρρητη.' )

    main_paragraph11 = doc.add_paragraph ( )
    main_paragraph11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph11.add_run ( 'Ζ.  Εκ του άρθρου 100 του Κ.Π.Δ., ήτοι δικαίωμα πρόσβασης στο υλικό της δικογραφίας:' )

    main_paragraph12 = doc.add_paragraph ( )
    main_paragraph12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph12.add_run ( 'Ανακοίνωση των εγγράφων της ανάκρισης. 1. Ο ανακριτής, μόλις μετά την κλήτευσή του '
                               'εμφανισθεί ή οδηγηθεί σε αυτόν ο κατηγορούμενος για να απολογηθεί, του ανακοινώνει το '
                               'περιεχόμενο του κατηγορητηρίου και των εγγράφων της ανάκρισης. Επιτρέπεται επίσης στον'
                               ' κατηγορούμενο να μελετήσει ο ίδιος ή ο συνήγορός του το κατηγορητήριο και τα έγγραφα '
                               'της ανάκρισης. Με γραπτή αίτηση του κατηγορουμένου και με δαπάνη του χορηγούνται σε '
                               'αυτόν αντίγραφα του κατηγορητηρίου και των εγγράφων της ανάκρισης. 2. Την ίδια υποχρέωση'
                               ' έχει ο ανακριτής, και τα ίδια δικαιώματα ο κατηγορούμενος, όταν κληθεί ξανά σε'
                               ' συμπληρωματική απολογία. Πάντως μετά το τέλος της ανάκρισης και προτού διαβιβαστεί η '
                               'δικογραφία στον εισαγγελέα (άρθρο 308 παρ. 1), καλείται πάντοτε ο κατηγορούμενος να '
                               'μελετήσει όλη τη δικογραφία. Αν όμως η ανάκριση εξακολούθησε περισσότερο από μήνα μετά'
                               ' την πρώτη ή κάθε μεταγενέστερη απολογία, δικαιούται ο κατηγορούμενος να ασκεί τα '
                               'δικαιώματά του μια φορά το μήνα, και κάθε φορά ο ανακριτής συντάσσει σχετική έκθεση κάτω'
                               ' από την απολογία του κατηγορουμένου.' )

    main_paragraph13 = doc.add_paragraph ( )
    main_paragraph13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph13.add_run ( 'Η.  Εκ του άρθρου 101 του Κ.Π.Δ. , ήτοι δικαίωμα διερμηνείας και μετάφρασης  :' )

    main_paragraph14 = doc.add_paragraph ( )
    main_paragraph14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph14.add_run (
        'Ο ύποπτος ή ο κατηγορούμενος, που δεν ομιλεί ή δεν κατανοεί επαρκώς την ελληνική γλώσσα, '
        'έχει το δικαίωμα σε διερμηνεία καθώς και σε γραπτή μετάφραση όλων των ουσιωδών εγγράφων '
        'της διαδικασίας, σύμφωνα με όσα προβλέπονται στα άρθρα 233 και 237. ' )

    main_paragraph15 = doc.add_paragraph ( )
    main_paragraph15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph15.add_run ( 'Θ. Εκ του άρθρου 103 του Κ.Π.Δ., ήτοι δικαίωμα προθεσμίας για την απολογία  :' )

    main_paragraph16 = doc.add_paragraph ( )
    main_paragraph16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph16.add_run (
        '1. Ο κατηγορούμενος έχει το δικαίωμα να ζητήσει προθεσμία τουλάχιστον σαράντα οκτώ ωρών '
        'και δεν έχει υποχρέωση να απολογηθεί πριν περάσει η προθεσμία. 2. Ο ανακριτής μπορεί να '
        'παρατείνει την προθεσμία ύστερα από αίτηση του κατηγορουμένου.  ' )

    main_paragraph17 = doc.add_paragraph ( )
    main_paragraph17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph17.add_run ( 'Ι.  Εκ του άρθρου 104 του Κ.Π.Δ. , ήτοι δικαίωμα σιωπής και μη αυτοενοχοποίησης :' )

    main_paragraph18 = doc.add_paragraph ( )
    main_paragraph18.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph18.add_run ( '1. Ο ύποπτος ή ο κατηγορούμενος έχουν δικαίωμα σιωπής και μη αυτοενοχοποίησης. '
                               '2. Η άσκηση του δικαιώματος μη αυτοενοχοποίησης δεν εμποδίζει τη νόμιμη συγκέντρωση '
                               'αποδεικτικών στοιχείων, που υπάρχουν ανεξάρτητα από τη βούληση των υπόπτων και των '
                               'κατηγορουμένων.'
                               ' 3. Η άσκηση του δικαιώματος σιωπής και μη αυτοενοχοποίησης δεν μπορεί να αξιοποιηθεί σε '
                               'βάρος των υπόπτων και των κατηγορουμένων. ' )

    main_paragraph19 = doc.add_paragraph ( )
    main_paragraph19.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph19.add_run ( 'ΙΑ.  Εκ του άρθρου 105 του Κ.Π.Δ. , ήτοι δικαιώματα στην αυτεπάγγελτη προανάκριση  :' )

    main_paragraph20 = doc.add_paragraph ( )
    main_paragraph20.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph20.add_run (
        '1.	Όταν ενεργείται προανάκριση σύμφωνα με το άρθρο 245 παρ. 2, η εξέταση γίνεται όπως '
        'ορίζεται στις διατάξεις των άρθρων 273 και 274 και εκείνος που εξετάζεται έχει τα '
        'δικαιώματα που αναφέρονται στα άρθρα 91, 95, 96, 97, 98, 100, 101, 103 και 104. '
        '2. Ο εξεταζόμενος έχει το δικαίωμα πρόσβασης σε δικηγόρο χωρίς αδικαιολόγητη καθυστέρηση '
        'και σε κάθε περίπτωση προτού εξεταστεί από την αστυνομία ή άλλη αρχή επιβολής του νόμου '
        'ή δικαστική αρχή. Η διάταξη του άρθρου 99 παρ. 4 εφαρμόζεται και σε αυτή την περίπτωση. '
        '3. Η κατά παράβαση του παρόντος άρθρου εξέταση είναι άκυρη και δεν λαμβάνεται υπόψη.'
        ' Κατά τα άλλα εφαρμόζεται η παρ. 3 του άρθρου 244.' )

    main_paragraph21 = doc.add_paragraph ( )
    main_paragraph21.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph21.add_run ( 'Κατόπιν της ανωτέρω γνωστοποίησης ο κατηγορούμενος δήλωσε ότι ενημερώθηκε'
                               ' πλήρως για τα παραπάνω δικαιώματα του και ερωτηθείς αν επιθυμεί να τα '
                               'ασκήσει απάντησε ότι δεν επιθυμεί να ασκήσει κάποιο από τα δικαιώματα '
                               'του ή ....... ' )

    main_paragraph22 = doc.add_paragraph ( )
    main_paragraph22.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph22.add_run ( 'Η παρούσα έκθεση άρχισε να συντάσσεται την {{start_time_rights1}}'
                               ' ώρα και περατώθηκε την {{end_time_rights1}}  ώρα' )

    main_paragraph23 = doc.add_paragraph ( )
    main_paragraph23.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση
    main_paragraph23.add_run ( 'Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού αναγνώσθηκε και βεβαιώθηκε,'
                               ' υπογράφεται ως ακολούθως:' )

    new_para5 = doc.add_paragraph ( "  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                    "         Ο   Ανακριτικός Υπάλληλος" )
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save ( 'sample_rights.docx' )
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ( "sample_rights.docx" )
    template.render ( context )
    template.save ( "ΔΙΚΑΙΩΜΑΤΑ_ΚΑΤΗΓΟΡΟΥΜΕΝΟΥ_mutual.docx" )
    # Αποθήκευση του εγγράφου




def generate_all():
    witness_report()
    iatrodik()
    silipsi()
    deltio_kat()
    apologia()
    rights()
    vas()
    transmission()


def witness_reportOfficerWeapon():
    global context
    doc = Document()
    title = doc.add_heading('ΕΚΘΕΣΗ ΕΝΟΡΚΗΣ ΕΞΕΤΑΣΗΣ (ΑΣΤΥΝΟΜΙΚΟΥ)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph = doc.add_paragraph()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run(
        '    Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}} ημέρα εβδομάδας '
        '{{day}} και ώρα {{hour}} ενώπιον εμού του {{first_officer}} του {{policeStation}} '
        'Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της ιδίας υπηρεσίας, που προσλήφθηκε '
        'ως Β\' Ανακριτικός Υπάλληλος, εμφανίστηκε ο κατωτέρω μάρτυρας, ο οποίος αφού ρωτήθηκε για την ταυτότητα '
        'του κ.λ.π. απάντησε ότι ονομάζεται: {{surname}} {{name}} του {{fathername}} και της {{mothername}} γεν. '
        '{{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου {{tel}}, ηλεκτρονικό '
        'ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}} '
        'Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}}'
    )
    new_para1 = doc.add_paragraph("   Έπειτα ο εξεταζόμενος έδωσε τον προβλεπόμενο από τα άρθρα 219 και 220 παρ. 1 "
                                  " του Κώδικα Ποινικής Δικονομίας όρκο, ως ακολούθως: «Δηλώνω, επικαλούμενος την "
                                  "τιμή και την συνείδηση μου, ότι θα πω όλη την αλήθεια και μόνο την αλήθεια, χωρίς"
                                  " να προσθέσω ούτε να αποκρύψω τίποτα», και στην συνέχεια εξετάσθηκε ως εξής:.")
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph("ΕΡΩΤΗΣΗ: Ρωτήθηκε σχετικά:")
    new_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para3 = doc.add_paragraph(
        "ΑΠΟΚΡΙΣΗ: Είμαι Αστυνομικός και υπηρετώ στο {{department}}. Στις {{ dateOfService}} βρισκόμουν σε "
        "διατεταγμένη Υπηρεσία {{kindOfService}} κατά τις ώρες {{hourOfService}}. Περί ώρα {{hourOfControl}}"
        "την {{dateOfControl}} ενώ βρισκόμασυταν με το περιπολικό στην περιοχή {{areaOfControl}} επί της οδού"
        "{{addressOfControl}} σταματήσαμε το υπ' αριθμόν {{vehicleNumber}} {{typeOfVehicle}}, μάρκας {{brand}}"
        " ιδιοκτησίας "
        "{{surnamePerperator}} {{namePerperator}} του {{fathernamePerperator}} και της {{mothernamePerperator}}"
        "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ.{{addressPerperator}},αριθμός τηλεφώνου "
        "{{telPreperator}}, ηλεκτρονικό ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}}"
        " που εκδόθηκε την {{issuedPerperator}} από {{place_issuedPerperator}} Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : "
        "{{doyPrep}} όπου κατά τον έλεγχο του οχήματος διαπιστώσαμε να έχει στo {{placeOfWeapon}} {{weapon}}. Ανωτέρω"
        " δράστης συνελήφθη, το παράνομο αντικείμενο κατασχέθηκε και οδηγήθηκε στο {{departmentOfDuty}}."
        "Τίποτε άλλο δεν έχω να προσθέσω και υπογράφω,")
    new_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para4 = doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hour}} ώρα και περαιώθηκε την "
                                  "{{hourOfReportFinished}} ώρα. Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού"
                                  "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5 = doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                  "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sample_document.docx')
    template = DocxTemplate("sample_document.docx")
    template.render(context)
    template.save("ΕνορκηΠερίΌπλων.docx")


def weapon_seizure():
    global context
    doc=Document()
    title=doc.add_heading('ΕΚΘΕΣΗ ΠΑΡΑΔΟΣΕΩΣ ΚΑΙ ΚΑΤΑΣΧΕΣΕΩΣ', level = 0)
    title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph=doc.add_paragraph()
    main_paragraph.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run( 'Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}}'
                            ' ημέρα εβδομάδας {{day}} και ώρα {{hourSeizure}} ενώπιον εμού του {{first_officer}} '
                            'του {{policeStation}} Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της '
                            'ιδίας υπηρεσίας, που προσλήφθηκε ως Β\' Ανακριτικός Υπάλληλος, επειδή ενεργούμε '
                            'προανάκριση για παράβαση του άρθρου 1 του Ν.2168/1993 ως αντικ. '
                            'με άρθρο 1 του Ν.4678/2020  (Περί όπλων) προβήκαμε στην κατάσχεσή του κάτωθι: '
                            '{{weapon}} που μας παρέδωσε ο/η {{surname}} {{name}} του {{fathername}} και της'
                            ' {{mothername}} γεν. {{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός '
                            'τηλεφώνου {{tel}}, ηλεκτρονικό ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DAT}} '
                            'που εκδόθηκε την {{issued}} από {{place_issued}} Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}}'
                            'το/τα οποίο/α βρέθηκαν στην κατοχή του/της {{surnamePerperator}} {{namePerperator}} του '
                            '{{fathernamePerperator}} και της {{mothernamePerperator}} γεν. {{dateOfBirthPerperator}} '
                            'στη {{placeOfBirthPerperator}} κατ.{{addressPerperator}},αριθμός τηλεφώνου '
                            '{{telPreperator}}, ηλεκτρονικό ταχυδρομείου {{emailPreperator}},κάτοχος του υπ αριθμόν '
                            '{{DATperperator}} που εκδόθηκε την {{issuedPerperator}} από {{place_issuedPerperator}} '
                            'Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : {{doyPrep}}, ύστερα από γενόμενο  αστυνομικό έλεγχο'
                            ' που ενεργήθηκε στις {{hourOfControl}} της {{dateOfControl}} στην περιοχή '
                            '{{areaOfControl}} επί της οδού {addressOfControl}}. Ανωτέρω αντικείμενο/α βρέθηκαν στο '
                            '{{placeOfWeapon}} στο υπ\' αριθμόν {{vehicleNumber}} {{typeOfVehicle}}, μάρκας {{brand}} '
                            'ιδιοκτησίας του'
                            'ανωτέρω.'
                            )
    new_para1=doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hourSeizure}} ώρα και περαιώθηκε την "
                                "{{hourOfSeizureFinished}} ώρα. Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία "
                                "αφού αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5=doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sample_document.docx')
    template=DocxTemplate("sample_document.docx")
    template.render(context)
    template.save("Εκθεση_Κατασχεσης.docx")



def witness_report_adjudication():
    global context
    doc = Document()
    title = doc.add_heading('ΕΚΘΕΣΗ ΕΝΟΡΚΗΣ ΕΞΕΤΑΣΗΣ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph = doc.add_paragraph()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run(
        '    Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}} ημέρα εβδομάδας '
        '{{day}} και ώρα {{hour}} ενώπιον εμού του {{first_officer}} του {{policeStation}} '
        'Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της ιδίας υπηρεσίας, που προσλήφθηκε '
        'ως Β\' Ανακριτικός Υπάλληλος, εμφανίστηκε ο κατωτέρω μάρτυρας, ο οποίος αφού ρωτήθηκε για την ταυτότητα '
        'του κ.λ.π. απάντησε ότι ονομάζεται: {{surname}} {{name}} του {{fathername}} και της {{mothername}} γεν. '
        '{{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου {{tel}}, ηλεκτρονικό '
        'ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}} '
        'Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}}'
    )
    new_para1 = doc.add_paragraph("   Έπειτα ο εξεταζόμενος έδωσε τον προβλεπόμενο από τα άρθρα 219 και 220 παρ. 1 "
                                  " του Κώδικα Ποινικής Δικονομίας όρκο, ως ακολούθως: «Δηλώνω, επικαλούμενος την "
                                  "τιμή και την συνείδηση μου, ότι θα πω όλη την αλήθεια και μόνο την αλήθεια, χωρίς"
                                  " να προσθέσω ούτε να αποκρύψω τίποτα», και στην συνέχεια εξετάσθηκε ως εξής:.")
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph("ΕΡΩΤΗΣΗ: Ρωτήθηκε σχετικά:")
    new_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para3 = doc.add_paragraph(
        "ΑΠΟΚΡΙΣΗ: Είμαι μόνιμος κάτοικος {{address}}. Από τις {{dateOfSeparation}} έχουμε χωρίσει με τον πρώην σύζυγο "
        "μου {{surnamePerperator}} {{namePerperator}} του {{fathernamePerperator}} και της {{mothernamePerperator}}"
        "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ.{{addressPerperator}},αριθμός τηλεφώνου "
        "{{telPreperator}}, ηλεκτρονικό ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}}"
        " που εκδόθηκε την {{issuedPerperator}} από {{place_issuedPerperator}} Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : "
        "{{doyPrep}} με τον οποίο έχω αποστήσει {{numberOfChildren}} {{nameOfChildren}}. Συνέπεια αυτού έχει συνταχθεί"
        "η υπ' αριθμόν {{numberOfAdjudication}} από {{dateOfAdjudication}} {{nameOfAdjudication}} όπου προβλέπει"
        "{{Adjudication}}. Στις {{dateOfCrime}} και περί ώρα {{hourOfCrime}} στη {{placeOfCrime}} ο πρώην σύζυγος μου"
        "παραβίασε την ανωτέρω δικαστική απόφαση {{whathappened}}.{{add_something}} Επιθυμώ την ποινική του δίωξη και"
        " σας εγχειρίζω την απόφαση.Τίποτε άλλο δεν έχω να προσθέσω και υπογράφω,")
    new_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para4 = doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hour}} ώρα και περαιώθηκε την "
                                  "{{hourOfReportFinished}} ώρα. Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού"
                                  "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5 = doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                  "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sample_document.docx')
    template = DocxTemplate("sample_document.docx")
    template.render(context)
    template.save("ΕνορκηΠαραβίαση Δικαστικής.docx")


def transmission_adjudication():
    global context  # Declare we're using the global context
    # Δημιουργία εγγράφου
    doc = Document()

    # Δημιουργία πίνακα με 1 γραμμή και 2 κελιά (για αριστερά & δεξιά κείμενο)
    table = doc.add_table(rows = 1,cols = 2)
    table.autofit = False  # Απενεργοποίηση αυτόματου προσαρμογής πλάτους

    # Ορισμός πλάτους στηλών (50% - 50%)
    col_widths = [Inches(3),Inches(3)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Κελί 1 (Αριστερά) - "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ"
    left_cell = table.cell(0, 0)
    left_texts = [
        "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ",
        "ΕΛΛΗΝΙΚΗ ΑΣΤΥΝΟΜΙΑ",
        "ΔΙΕΥΘ. ΑΣΤ. ΘΕΣΣΑΛΟΝΙΚΗΣ",
        "ΑΣΤΥΝΟΜΙΚΟ ΤΜΗΜΑ ΘΕΡΜΗΣ",
        "ΤΗΛ:2310461203",
        "email: atthermis@astynomia.gr",
        "Αρμόδιος:{{first_officer }}",
        "Αρ.πρωτ: {{ protocolnumber }}"
    ]

    for text in left_texts:
        p = left_cell.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)  # γραμματοσειρά (12pt)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)

    # Κελί 2 (Δεξιά) - Τοποθεσία, ημερομηνία + επιπλέον κείμενο
    right_cell = table.cell ( 0 ,1 )

    # Γραμμή 1: Τοποθεσία & Ημερομηνία
    right_para1 = right_cell.add_paragraph ( )
    right_run1 = right_para1.add_run ( f"Θέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
    right_run1.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Γραμμή 2: "Π Ρ Ο Σ" (με κενά)
    right_para2 = right_cell.add_paragraph ( )
    right_run2 = right_para2.add_run ( " Π Ρ Ο Σ" )
    right_run2.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Γραμμή 3: Εργαστήριο & Πανεπιστήμιο
    right_para3 = right_cell.add_paragraph ( )
    right_run3 = right_para3.add_run (
        "ΕΙΣΑΓΓΕΛΙΑ ΠΛΗΜ/ΚΩΝ \nΘΕΣΣΑΛΟΝΙΚΗΣ\n" )
    right_run3.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph ( ).add_run ( ).add_break ( )  # Κενή γραμμή πριν το ΘΕΜΑ

    p_theme = doc.add_paragraph ( )
    run_theme = p_theme.add_run ( "ΘΕΜΑ: «Αποστολή Συνοδείας»" )
    run_theme.bold = True
    run_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    p_theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν

    main_theme = doc.add_paragraph ( )
    run1_theme = main_theme.add_run ( ' Υποβάλλεται συννημένα δικογραφία που σχηματίσθηκε στην Υπηρεσία μας και αφορά  '
                                      'την υποβολή έγκλησης της  {{surname}} {{name}} του {{fathername}} και της '
                                      '{{mothername}} γεν.{{dateOfBirth}} στη {{placeOfBirth}} κατ.{{address}}, αριθμός'
                                      'τηλεφώνου {{tel}}, ηλεκτρονικό ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν '
                                      '{{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}} Α.Φ.Μ : {{afm}}, Δ.Ο.Υ '
                                      ': {{doy}} εις βάρος του {surnamePerperator}} {{namePerperator}} του '
                                      '{{fathernamePerperator}} και της {{mothernamePerperator}}'
                                      'γεν. {{dateOfBirthPerperator }} στη {{placeOfBirthPerperator }} κατ.{'
                                      '{ addressPerperator }},αριθμός τηλεφώνου {{ telPreperator }}, ηλεκτρονικό '
                                      'ταχυδρομείου{{ emailPreperator }}, κάτοχος του υπ αριθμόν {{ DATperperator}} που'
                                      ' εκδόθηκε την {{ issuedPerperator }} από {{place_issuedPerperator}} Α.Φ.Μ : '
                                      '{{afmPreperator}}, Δ.Ο.Υ : {{ doyPrep }}, για παράβαση του/των άρθρου/ων '
                                      '{{offences}} πράξη/εις που έλαβαν χώρα  στη {{placeOfCrime}} στις '
                                      '{{ dateOfCrime}} και περί ώρα {{ hourOfCrime }}.\n' 
                                      '\tΣυγκεκριμένα ανωτέρω τόπο και χρόνο δράστης παραβίασε την υπ αριθμόν'
                                      '{{numberOfAdjudication}} από {{dateOfAdjudication}} {{nameOfAdjudication}} όπου '
                                      'προβλέπει {{Adjudication}}'
                                      )
    run1_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση

    footer_para = doc.add_paragraph ( )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Προσθήκη κειμένου με ακριβή θέση (χρήση tabs)
    footer_run = footer_para.add_run ( )
    footer_run.add_text(f"\t\tΘέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
    footer_run.add_text("\t\t\tΟ\t\t\n")  # 5 tabs για στοίχιση
    footer_run.add_text("\t\t\t\t\tΑνακριτικός Υπάλληλος\n\n")
    footer_run.add_text("\t\t\t\t\t {{ first_officer }}\n")

    # Αποθήκευση
    doc.save('sample_transmission.docx')
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ('sample_transmission.docx')
    template.render(context)
    template.save("ΔΙΑΒΙΒΑΣΤΙΚΟ_παραβίαση_δικαστικης.docx")
    # Αποθήκευση του εγγράφου



def transmission_gunpos():
    global context  # Declare we're using the global context
    # Δημιουργία εγγράφου
    doc = Document()

    # Δημιουργία πίνακα με 1 γραμμή και 2 κελιά (για αριστερά & δεξιά κείμενο)
    table = doc.add_table(rows = 1,cols = 2)
    table.autofit = False  # Απενεργοποίηση αυτόματου προσαρμογής πλάτους

    # Ορισμός πλάτους στηλών (50% - 50%)
    col_widths = [Inches(3),Inches(3)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Κελί 1 (Αριστερά) - "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ"
    left_cell = table.cell(0, 0)
    left_texts = [
        "ΕΛΛΗΝΙΚΗ ΔΗΜΟΚΡΑΤΙΑ",
        "ΕΛΛΗΝΙΚΗ ΑΣΤΥΝΟΜΙΑ",
        "ΔΙΕΥΘ. ΑΣΤ. ΘΕΣΣΑΛΟΝΙΚΗΣ",
        "ΑΣΤΥΝΟΜΙΚΟ ΤΜΗΜΑ ΘΕΡΜΗΣ",
        "ΤΗΛ:2310461203",
        "email: atthermis@astynomia.gr",
        "Αρμόδιος:{{first_officer }}",
        "Αρ.πρωτ: {{ protocolnumber }}"
    ]

    for text in left_texts:
        p = left_cell.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)  # γραμματοσειρά (12pt)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)

    # Κελί 2 (Δεξιά) - Τοποθεσία, ημερομηνία + επιπλέον κείμενο
    right_cell = table.cell ( 0 ,1 )

    # Γραμμή 1: Τοποθεσία & Ημερομηνία
    right_para1 = right_cell.add_paragraph ( )
    right_run1 = right_para1.add_run ( f"Θέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
    right_run1.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Γραμμή 2: "Π Ρ Ο Σ" (με κενά)
    right_para2 = right_cell.add_paragraph ( )
    right_run2 = right_para2.add_run ( " Π Ρ Ο Σ" )
    right_run2.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Γραμμή 3: Εργαστήριο & Πανεπιστήμιο
    right_para3 = right_cell.add_paragraph ( )
    right_run3 = right_para3.add_run (
        "ΕΙΣΑΓΓΕΛΙΑ ΠΛΗΜ/ΚΩΝ \nΘΕΣΣΑΛΟΝΙΚΗΣ\n" )
    right_run3.font.size = Pt ( 12 )  # γραμματοσειρά (12pt)
    right_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph ( ).add_run ( ).add_break ( )  # Κενή γραμμή πριν το ΘΕΜΑ

    p_theme = doc.add_paragraph ( )
    run_theme = p_theme.add_run ( "ΘΕΜΑ: «Αποστολή Συνοδείας»" )
    run_theme.bold = True
    run_theme.font.size = Pt ( 14 )  # γραμματοσειρά (14pt)
    p_theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_theme.paragraph_format.space_before = Pt ( 16 )  # Επιπλέον κενό πριν

    main_theme = doc.add_paragraph ( )
    run1_theme = main_theme.add_run ( ' Υποβάλλεται συννημένα δικογραφία που σχηματίσθηκε στην Υπηρεσία μας και αφορά  '
                                      'την υποβολή έγκλησης της  {{surname}} {{name}} του {{fathername}} και της '
                                      '{{mothername}} γεν.{{dateOfBirth}} στη {{placeOfBirth}} κατ.{{address}}, αριθμός'
                                      'τηλεφώνου {{tel}}, ηλεκτρονικό ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν '
                                      '{{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}} Α.Φ.Μ : {{afm}}, Δ.Ο.Υ '
                                      ': {{doy}} εις βάρος του {surnamePerperator}} {{namePerperator}} του '
                                      '{{fathernamePerperator}} και της {{mothernamePerperator}}'
                                      'γεν. {{dateOfBirthPerperator }} στη {{placeOfBirthPerperator }} κατ.{'
                                      '{ addressPerperator }},αριθμός τηλεφώνου {{ telPreperator }}, ηλεκτρονικό '
                                      'ταχυδρομείου{{ emailPreperator }}, κάτοχος του υπ αριθμόν {{ DATperperator}} που'
                                      ' εκδόθηκε την {{ issuedPerperator }} από {{place_issuedPerperator}} Α.Φ.Μ : '
                                      '{{afmPreperator}}, Δ.Ο.Υ : {{ doyPrep }}, για παράβαση της Νομοθεσία Περί Όπλων '
                                      'και Εκρηκτικών - 2168/93'
                                      ' πράξη/εις που έλαβαν χώρα  στη {{areaOfControl}} στις '
                                      '{{dateOfControl}} και περί ώρα {{ hourOfControl }}.\n' 
                                      '\tΣυγκεκριμένα ανωτέρω τόπο και χρόνο δράστης διαπιστώθηκε ύστερα από Αστυνομικό'
                                      'έλεγχο που διενεργήθηκε στο υπ αριθμόν {{vehicleNumber}} {{typeOfVehicle}} ' 
                                      'μάρκας {{brand}} ιδιοκτησίας του, να κατέχει εντός αυτού και συγκεκριμένα στο'
                                      '{{placeOfWeapon}} {{weapon}}. Ανωτέρω αντικείμενο κατασχέθηκε.'
                                      '  Παρακαλούμε για τις δικές σας ενέργειες.'

                                      )
    run1_theme.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Εδώ ορίζουμε την πλήρη στοίχιση

    footer_para = doc.add_paragraph ( )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Προσθήκη κειμένου με ακριβή θέση (χρήση tabs)
    footer_run = footer_para.add_run ( )
    footer_run.add_text(f"\t\tΘέρμη, {datetime.now ( ).strftime ( '%d/%m/%Y' )}" )
    footer_run.add_text("\t\t\tΟ\t\t\n")  # 5 tabs για στοίχιση
    footer_run.add_text("\t\t\t\t\tΑνακριτικός Υπάλληλος\n\n")
    footer_run.add_text("\t\t\t\t\t {{ first_officer }}\n")

    # Αποθήκευση
    doc.save('sample_transmission.docx')
    # Επεξεργασία και τελική αποθήκευση
    template = DocxTemplate ('sample_transmission.docx')
    template.render(context)
    template.save("ΔΙΑΒΙΒΑΣΤΙΚΟ_περι_όπλων.docx")
    # Αποθήκευση του εγγράφου


def witness_report_oblig():
    global context
    doc = Document()
    title = doc.add_heading('ΕΚΘΕΣΗ ΕΝΟΡΚΗΣ ΕΞΕΤΑΣΗΣ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_paragraph = doc.add_paragraph()
    main_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_paragraph.add_run(
        '    Στην {{place}} σήμερα την {{date_num}} του μήνα {{month}} του έτους {{year}} ημέρα εβδομάδας '
        '{{day}} και ώρα {{hour}} ενώπιον εμού του {{first_officer}} του {{policeStation}} '
        'Θεσσαλονίκης, παρισταμένου και της {{sec_officer}} της ιδίας υπηρεσίας, που προσλήφθηκε '
        'ως Β\' Ανακριτικός Υπάλληλος, εμφανίστηκε ο κατωτέρω μάρτυρας, ο οποίος αφού ρωτήθηκε για την ταυτότητα '
        'του κ.λ.π. απάντησε ότι ονομάζεται: {{surname}} {{name}} του {{fathername}} και της {{mothername}} γεν. '
        '{{dateOfBirth}} στη {{placeOfBirth}} κατ. {{address}}, αριθμός τηλεφώνου {{tel}}, ηλεκτρονικό '
        'ταχυδρομείο {{email}}, κάτοχος του υπ αριθμόν {{DAT}} που εκδόθηκε την {{issued}} από {{place_issued}} '
        'Α.Φ.Μ : {{afm}}, Δ.Ο.Υ : {{doy}}'
    )
    new_para1 = doc.add_paragraph("   Έπειτα ο εξεταζόμενος έδωσε τον προβλεπόμενο από τα άρθρα 219 και 220 παρ. 1 "
                                  " του Κώδικα Ποινικής Δικονομίας όρκο, ως ακολούθως: «Δηλώνω, επικαλούμενος την "
                                  "τιμή και την συνείδηση μου, ότι θα πω όλη την αλήθεια και μόνο την αλήθεια, χωρίς"
                                  " να προσθέσω ούτε να αποκρύψω τίποτα», και στην συνέχεια εξετάσθηκε ως εξής:.")
    new_para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para2 = doc.add_paragraph("ΕΡΩΤΗΣΗ: Ρωτήθηκε σχετικά:")
    new_para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para3 = doc.add_paragraph(
        "ΑΠΟΚΡΙΣΗ: Είμαι μόνιμος κάτοικος {{address}}. Από τις {{dateOfSeparation}} έχουμε χωρίσει με τον πρώην σύζυγο "
        "μου {{surnamePerperator}} {{namePerperator}} του {{fathernamePerperator}} και της {{mothernamePerperator}}"
        "γεν. {{dateOfBirthPerperator}} στη {{placeOfBirthPerperator}} κατ.{{addressPerperator}},αριθμός τηλεφώνου "
        "{{telPreperator}}, ηλεκτρονικό ταχυδρομείου{{emailPreperator}},κάτοχος του υπ αριθμόν {{DATperperator}}"
        " που εκδόθηκε την {{issuedPerperator}} από {{place_issuedPerperator}} Α.Φ.Μ : {{afmPreperator}}, Δ.Ο.Υ : "
        "{{doyPrep}} με τον οποίο έχω αποστήσει {{numberOfChildren}} {{nameOfChildren}}. Συνέπεια αυτού έχει συνταχθεί"
        "η υπ' αριθμόν {{numberOfAdjudication}} από {{dateOfAdjudication}} {{nameOfAdjudication}} όπου προβλέπει"
        "{{obligation}}. Στις {{dateOfCrime}} και περί ώρα {{hourOfCrime}} στη {{placeOfCrime}} ο πρώην σύζυγος μου"
        "παραβίασε την ανωτέρω δικαστική απόφαση {{whathappened}}.{{add_something}} Επιθυμώ την ποινική του δίωξη και"
        " σας εγχειρίζω την απόφαση.Τίποτε άλλο δεν έχω να προσθέσω και υπογράφω,")
    new_para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para4 = doc.add_paragraph("Η παρούσα έκθεση άρχισε να συντάσσεται την {{hour}} ώρα και περαιώθηκε την "
                                  "{{hourOfReportFinished}} ώρα. Για πίστωση συντάχθηκε η παρούσα έκθεση η οποία αφού"
                                  "αναγνώσθηκε και βεβαιώθηκε, υπογράφεται ως ακολούθως:")
    new_para4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_para5 = doc.add_paragraph("  Ο Εξετασθείς                       Ο Β Ανακριτικός Υπάλληλος       "
                                  "         Ο   Ανακριτικός Υπάλληλος")
    new_para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save('sample_document.docx')
    template = DocxTemplate("sample_document.docx")
    template.render(context)
    template.save("ΕνορκηΠαραβίαση Δικαστικής.docx")


